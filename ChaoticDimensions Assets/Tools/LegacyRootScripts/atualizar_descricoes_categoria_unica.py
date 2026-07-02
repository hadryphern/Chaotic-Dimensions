from __future__ import annotations

import re
import shutil
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from gerar_relatorio_completo_docx import add_hyperlink, clean, strip_html
from gerar_relatorio_humanizado_separado import (
    add_supplier_placeholder_issues,
    load_review_products,
    remove_deductive_issues,
    reviewed_issues,
)


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "relatorio_clickmed_HUMANIZADO_separado.docx"
LATEST_DOCX = ROOT / "relatorio_clickmed_FINAL_REVISADO.docx"
FINAL_BLOCK_DOCX = ROOT / "relatorio_clickmed_final_blocos.docx"
BACKUP_DIR = ROOT / "backups_clickmed"
VALIDATION = ROOT / "validacao_relatorio_clickmed_descricoes_categoria_unica.md"

DESCRIPTION_KINDS = {
    "Descricao/IA",
    "Descricao/confirmar fornecedor",
    "Descricao/avaliacao",
    "Descricao/campos vazios",
    "Descricao/imagem externa",
    "Descricao",
    "Descricao/relacionados texto",
}


def add_related_plain_text_issues(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, list[dict]]:
    enriched = {product_id: list(items) for product_id, items in issues.items()}
    for product in products:
        raw = (product.get("short_description") or "") + "\n" + (product.get("description") or "")
        raw_lower = raw.lower()
        marker_index = raw_lower.find("pesquisas relacionadas")
        if marker_index < 0:
            marker_index = raw_lower.find("relacionados")
        if marker_index < 0:
            continue

        next_heading = raw_lower.find("<h2", marker_index + 10)
        section = raw_lower[marker_index: next_heading if next_heading > marker_index else len(raw_lower)]
        if "<li" not in section or "<a " in section or "<a\t" in section:
            continue

        item = {
            "kind": "Descricao/relacionados texto",
            "detail": "Pesquisas relacionadas/tags de pesquisa aparecem como texto simples, sem link clicavel.",
            "expected": "Transformar esses termos em links uteis ou remover a secao se for apenas enchimento de SEO.",
        }
        enriched.setdefault(product["id"], [])
        if item not in enriched[product["id"]]:
            enriched[product["id"]].append(item)
    return enriched


def description_issues_by_product(products: list[dict]) -> dict[int, list[dict]]:
    issues = reviewed_issues(products)
    issues = remove_deductive_issues(issues)
    issues = add_supplier_placeholder_issues(products, issues)
    issues = add_related_plain_text_issues(products, issues)

    desc_only: dict[int, list[dict]] = {}
    for product_id, items in issues.items():
        kept = [item for item in items if item["kind"] in DESCRIPTION_KINDS]
        if kept:
            desc_only[product_id] = kept
    return desc_only


def short_issue(item: dict) -> str:
    kind = item["kind"]
    detail = item.get("detail", "")
    if kind == "Descricao/confirmar fornecedor":
        examples = detail.replace("Descricao mostra placeholder interno publicado:", "").strip()
        return f"campos internos publicados como confirmar fornecedor/vendedor ({examples})"
    if kind == "Descricao/IA":
        return "problemas de IA/texto colado na descricao"
    if kind == "Descricao/avaliacao":
        return "avaliacao generica/4.8 sem reviews reais no produto"
    if kind == "Descricao/campos vazios":
        return "campos vazios na descricao"
    if kind == "Descricao/imagem externa":
        return "imagem externa dentro da descricao"
    if kind == "Descricao/relacionados texto":
        return "pesquisas relacionadas/tags em texto simples, sem link clicavel"
    if kind == "Descricao":
        return detail or "descricao ausente ou texto/template problemático"
    return detail


def short_fix(item: dict) -> str:
    kind = item["kind"]
    if kind == "Descricao/confirmar fornecedor":
        return "preencher com informacao real ou remover esses campos"
    if kind == "Descricao/IA":
        return "limpar restos de IA e revisar o texto"
    if kind == "Descricao/avaliacao":
        return "remover avaliacao generica ou usar reviews reais"
    if kind == "Descricao/campos vazios":
        return "preencher ou remover campos vazios"
    if kind == "Descricao/imagem externa":
        return "trocar por imagem propria/hospedada no site ou remover"
    if kind == "Descricao/relacionados texto":
        return "transformar termos relacionados em links uteis ou remover a secao"
    if kind == "Descricao":
        return item.get("expected") or "revisar a descricao"
    return item.get("expected") or "revisar manualmente"


def add_big_heading(doc: Document, text: str, level: int = 2, size: float = 22):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 105, 0)
    return paragraph


def set_size(paragraph, size: float):
    for run in paragraph.runs:
        run.font.size = Pt(size)


def add_note(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)
    set_size(paragraph, 9.5)


def add_link_line(doc: Document, url: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run("Link: ").bold = True
    add_hyperlink(paragraph, url, url)
    set_size(paragraph, 8.8)


def add_labeled_line(doc: Document, label: str, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run(label).bold = True
    paragraph.add_run(text)
    set_size(paragraph, 9.2)


def remove_from(doc: Document, start_index: int):
    for paragraph in list(doc.paragraphs[start_index:]):
        element = paragraph._element
        element.getparent().remove(element)


def find_description_start(doc: Document) -> int:
    markers = [
        "parte 2",
        "descricao/ia",
        "início da categoria: restos de ia",
        "inicio da categoria: restos de ia",
        "campos com confirmar fornecedor",
        "descricao/confirmar fornecedor",
    ]
    candidates = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if any(marker in text for marker in markers):
            candidates.append(index)
    if not candidates:
        raise RuntimeError("Nao encontrei onde comeca a area de descricoes no DOCX.")
    return min(candidates)


def build_unified_description_section(doc: Document, products: list[dict], issues: dict[int, list[dict]]):
    products_by_id = {product["id"]: product for product in products}
    ordered = sorted(
        ((products_by_id[product_id], items) for product_id, items in issues.items()),
        key=lambda row: (clean(row[0].get("name")).lower(), row[0]["id"]),
    )

    doc.add_page_break()
    add_big_heading(doc, f"Analise das descricoes dos produtos ({len(ordered)})", level=2, size=23)
    add_note(
        doc,
        "Esta parte ficou toda numa categoria unica para o colega revisar texto sem ter que abrir varias subcategorias.",
    )
    add_note(
        doc,
        "Em cada item, juntei os problemas da descricao na mesma linha: confirmar fornecedor/vendedor, restos de IA, avaliacao generica, campos vazios, imagem externa e termos relacionados sem link.",
    )

    counts = defaultdict(int)
    for items in issues.values():
        for item in items:
            counts[item["kind"]] += 1
    summary = "; ".join(
        [
            f"confirmar fornecedor/vendedor: {counts['Descricao/confirmar fornecedor']}",
            f"IA/texto colado: {counts['Descricao/IA']}",
            f"avaliacao generica: {counts['Descricao/avaliacao']}",
            f"campos vazios: {counts['Descricao/campos vazios']}",
            f"imagem externa: {counts['Descricao/imagem externa']}",
            f"relacionados/tags sem link: {counts['Descricao/relacionados texto']}",
            f"descricao ausente/template ruim: {counts['Descricao']}",
        ]
    )
    add_labeled_line(doc, "Resumo da categoria: ", summary)

    for product, items in ordered:
        title = doc.add_paragraph()
        title.paragraph_format.space_before = Pt(8)
        title.paragraph_format.space_after = Pt(2)
        title.add_run(f"{clean(product.get('name'))} | ID {product['id']}").bold = True
        set_size(title, 10)

        add_link_line(doc, product["permalink"])
        description_line = "; ".join(short_issue(item) for item in items)
        fix_line = "; ".join(dict.fromkeys(short_fix(item) for item in items))
        add_labeled_line(doc, "Descricao: ", description_line)
        add_labeled_line(doc, "Corrigir: ", fix_line)


def validate(docx: Path, products: list[dict], issues: dict[int, list[dict]], backup: Path):
    doc = Document(docx)
    with zipfile.ZipFile(docx) as archive:
        bad = archive.testzip()
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs).lower()
    lines = [
        "# Validacao - descricoes em categoria unica",
        "",
        f"- DOCX integro: {'sim' if bad is None else 'nao'}",
        f"- Backup criado: {backup.name}",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com problema de descricao: {len(issues)}",
        f"- Total de problemas de descricao: {sum(len(items) for items in issues.values())}",
        f"- Tabelas no DOCX: {len(doc.tables)}",
        f"- Imagens no DOCX: {len(doc.inline_shapes)}",
        f"- Ocorrencias de 'DESCRICAO/IA': {text.count('descricao/ia')}",
        f"- Ocorrencias de 'DESCRICAO/AVALIACAO': {text.count('descricao/avaliacao')}",
        f"- Ocorrencias de 'INICIO DA CATEGORIA: CAMPOS COM CONFIRMAR': {text.count('inicio da categoria: campos com confirmar')}",
        f"- Ocorrencias de 'pesquisas relacionadas/tags': {text.count('pesquisas relacionadas/tags')}",
    ]
    VALIDATION.write_text("\n".join(lines), encoding="utf-8")


def main():
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}_antes_descricoes_categoria_unica_{stamp}.docx"
    shutil.copyfile(DOCX, backup)

    products = load_review_products()
    desc_issues = description_issues_by_product(products)

    doc = Document(DOCX)
    start = find_description_start(doc)
    remove_from(doc, start)
    build_unified_description_section(doc, products, desc_issues)
    doc.save(DOCX)

    shutil.copyfile(DOCX, LATEST_DOCX)
    shutil.copyfile(DOCX, FINAL_BLOCK_DOCX)
    validate(DOCX, products, desc_issues, backup)

    print(DOCX)
    print(backup)
    print(VALIDATION)


if __name__ == "__main__":
    main()
