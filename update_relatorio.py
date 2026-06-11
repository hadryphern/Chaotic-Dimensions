# -*- coding: utf-8 -*-
"""
Insere registos de desenvolvimento no RELATÓRIO_CHAOTIC.docx.
Uso único — executar uma vez para aplicar as entradas da sessão de 11/06/2026.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

RELATORIO = os.path.join(os.path.dirname(__file__), "RELATÓRIO_CHAOTIC.docx")
doc = Document(RELATORIO)

# ── Helpers de inserção ───────────────────────────────────────────────────────

def _fmt_run(run, bold=False):
    run.font.name  = "Arial"
    run.font.size  = Pt(11)
    run.font.bold  = bold

def _fmt_para(p, first_indent=True):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(12)
    if first_indent:
        pf.first_line_indent = Cm(0.6)
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:line"),     "360")
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:after"),    "240")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(sp)

def insert_before(ref_elem, *new_elems):
    """Insere elementos XML imediatamente antes de ref_elem."""
    for e in reversed(new_elems):
        ref_elem.addprevious(e)

def build_h2(doc, text):
    """Cria um parágrafo Heading 2 temporariamente no fim e devolve o elemento."""
    p = doc.add_heading(text, level=2)
    elem = p._element
    doc.element.body.remove(elem)
    return elem

def build_normal(doc, text, bold_prefix=None):
    """Cria parágrafo Normal temporariamente no fim e devolve o elemento."""
    p = doc.add_paragraph()
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        _fmt_run(r0, bold=True)
    r1 = p.add_run(text)
    _fmt_run(r1)
    p.style = doc.styles["Normal"]
    _fmt_para(p)
    elem = p._element
    doc.element.body.remove(elem)
    return elem

def build_table_row_3(doc, c1, c2, c3):
    """Cria linha de tabela com 3 células (para tabela já existente)."""
    from docx.oxml import OxmlElement
    tr = OxmlElement('w:tr')
    for text in [c1, c2, c3]:
        tc   = OxmlElement('w:tc')
        p    = OxmlElement('w:p')
        r    = OxmlElement('w:r')
        rpr  = OxmlElement('w:rPr')
        sz   = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')
        rpr.append(sz)
        r.append(rpr)
        t    = OxmlElement('w:t'); t.text = text; t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
    return tr

# ── 1. Localizar âncoras no documento ────────────────────────────────────────

para_map = {}  # label → paragraph index
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    sn = p.style.name
    if sn == "Heading 1" and "Conteúdo Atual" in t:
        para_map["conteudo_atual"] = i
    if sn == "Heading 1" and "Melhorias" in t:
        para_map["melhorias"] = i
    if sn == "Heading 2" and "Evolução de versões" in t:
        para_map["evolucao"] = i
    if sn == "Heading 1" and "Reflexão" in t:
        para_map["reflexao"] = i

print("Âncoras encontradas:", para_map)

# ── 2. Inserir nova subsecção em Desenvolvimento ──────────────────────────────
# Inserir antes de "Conteúdo Atual" (que é logo após o fim do Desenvolvimento)

ref_conteudo = doc.paragraphs[para_map["conteudo_atual"]]._element

els_dev = [
    build_h2(doc, "Boss Mago Imã — Conceção e Spritesheet Idle"),
    build_normal(doc, (
        "A 11 de junho de 2026, deu-se início ao desenvolvimento do boss Mago Imã "
        "(denominado internamente MagnetMage), o primeiro dos bosses ChaoticApex "
        "a receber implementação de arte. O boss foi concebido como um mago sem pernas "
        "cujo corpo é coberto por um manto preto ao estilo detetive, com uma cabeça em "
        "forma de ímão de ferradura (corpo vermelho com pontas cinzentas), braços e mãos "
        "flutuantes desconectados do tronco e luvas de cor cinzento-escuro."
    )),
    build_normal(doc, (
        "Foi gerada a spritesheet de animação idle através de um script Python "
        "(gen_magnet_mage.py) utilizando as bibliotecas Pillow e NumPy, com supersampling "
        "de 4× para obter suavização de bordas equivalente a anti-aliasing. A spritesheet "
        "resultante (MagnetMageBoss.png) contém 120 fotogramas organizados em atlas "
        "de 10 colunas por 12 linhas, com cada fotograma a medir 220 × 300 píxeis, "
        "totalizando uma resolução de 2200 × 3600 píxeis."
    )),
    build_normal(doc, (
        "A animação idle inclui quatro elementos sincronizados: flutuação vertical do "
        "corpo (amplitude de 7 píxeis, ciclo completo de 2 segundos a 60 fps); balanço "
        "do manto com duas frequências ligeiramente distintas para cada lado, criando "
        "movimento orgânico; posição das mãos com desfasamento temporal de 14% do ciclo "
        "em relação ao corpo; e animação de abertura e fecho dos dedos em dois ciclos por "
        "período de flutuação, variando entre punho fechado e mão aberta."
    )),
    build_normal(doc, (
        "Os ficheiros criados foram: "
        "Content/NPCs/MagnetMage/MagnetMageBoss.png (spritesheet final), "
        "Content/NPCs/MagnetMage/MagnetMageBoss_preview.png (frame 0 em resolução 4× "
        "para revisão visual) e gen_magnet_mage.py (script gerador reutilizável na raiz "
        "do projeto). A spritesheet está preparada para integração com o sistema de atlas "
        "já utilizado pelo Alien Kraken, devendo o código C# do boss referenciar as "
        "constantes ATLAS_COLS = 10 e ATLAS_ROWS = 12."
    )),
]

insert_before(ref_conteudo, *reversed(els_dev))
print("✓ Subsecção 'Boss Mago Imã' inserida em Desenvolvimento")

# ── 3. Inserir erros na tabela de Melhorias ──────────────────────────────────
# Procurar a tabela seguinte ao heading "Melhorias e Correção de Erros"
melhoria_heading = doc.paragraphs[para_map["melhorias"]]._element

# Encontrar a primeira tabela após esse heading
tables_after = []
for tbl in doc.tables:
    tbl_pos = tbl._element.getparent().index(tbl._element)
    head_pos = tbl._element.getparent().index(melhoria_heading)
    if tbl_pos > head_pos:
        tables_after.append(tbl)

if tables_after:
    melhoria_tbl = tables_after[0]
    erros = [
        ("PIL/Pillow não instalado (ModuleNotFoundError)",
         "Módulo Pillow não estava presente no ambiente Python 3.13",
         "Instalado com python -m pip install Pillow"),
        ("NumPy não importado a nível de módulo (ModuleNotFoundError)",
         "Import de numpy dentro da função draw_frame em vez de no topo do ficheiro",
         "Movido import numpy as np para o cabeçalho do script"),
        ("Efeitos de brilho/glow com raio excessivo (background vermelho/roxo)",
         "Factores de escala dos glows (até 2.8×) geravam elipses de 560 px no "
         "canvas de trabalho 4×, preenchendo o fundo inteiro após GaussianBlur",
         "Raios reduzidos para máximo de 46 px (1× final), blur reduzido de s(7) para s(3)"),
    ]
    for c1, c2, c3 in erros:
        tr = build_table_row_3(doc, c1, c2, c3)
        melhoria_tbl._element.append(tr)
    print(f"✓ {len(erros)} erros adicionados à tabela de Melhorias")
else:
    # Inserir parágrafo se não houver tabela
    ref_evol = doc.paragraphs[para_map.get("evolucao", para_map["melhorias"] + 1)]._element
    note = build_normal(doc, (
        "11/06/2026 — Erros no script gen_magnet_mage.py: Pillow não instalado "
        "(resolvido com pip install Pillow); numpy importado dentro da função "
        "(movido para topo); glow com raio excessivo preenchia o fundo (raios reduzidos)."
    ))
    insert_before(ref_evol, note)
    print("✓ Nota de erros inserida (sem tabela encontrada)")

# ── 4. Guardar ────────────────────────────────────────────────────────────────
doc.save(RELATORIO)
print(f"\n✓ RELATÓRIO_CHAOTIC.docx actualizado e guardado.")
