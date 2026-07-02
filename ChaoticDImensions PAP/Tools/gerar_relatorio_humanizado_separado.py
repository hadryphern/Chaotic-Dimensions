from __future__ import annotations

import re
import shutil
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gerar_relatorio_blocos_final_docx import (
    REVIEW_REMOVED,
    confirmed_empty_categories,
    confirmed_security_findings,
    load_review_products,
    reviewed_issues,
)
from gerar_relatorio_completo_docx import add_hyperlink, clean, product_brands, product_categories, product_tags, strip_html


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_HUMANIZADO_separado.docx"
OUT_MD = ROOT / "relatorio_clickmed_HUMANIZADO_separado.md"
OUT_VALIDATION = ROOT / "validacao_relatorio_clickmed_humanizado_separado.md"

LATEST_DOCX = ROOT / "relatorio_clickmed_FINAL_REVISADO.docx"
LATEST_MD = ROOT / "relatorio_clickmed_FINAL_REVISADO.md"
FINAL_BLOCK_DOCX = ROOT / "relatorio_clickmed_final_blocos.docx"
FINAL_BLOCK_MD = ROOT / "relatorio_clickmed_final_blocos.md"

DESCRIPTION_KINDS = {
    "Descricao/IA",
    "Descricao/confirmar fornecedor",
    "Descricao/avaliacao",
    "Descricao/campos vazios",
    "Descricao/imagem externa",
    "Descricao",
}

DEDUCTIVE_KINDS = {
    "Tag/condicao",
    "Tag/Grade",
    "Tags",
    "Categoria",
    "Padronizacao",
    "Descricao/especificacao",
    "Descricao/Grade",
}

FILTERED_DEDUCTIONS: list[tuple[int, str, str]] = []

BOSS_GROUPS = [
    (
        "Marca, icones e filtros",
        "Aqui entram os casos em que a marca ou o filtro do produto pode levar o cliente para o caminho errado.",
        ["Marca"],
    ),
    (
        "Tags, estado e grade",
        "Esta parte e importante porque mexe diretamente com a expectativa do cliente: novo, usado, open box ou Grade A/B/C.",
        ["Tag/condicao", "Tag/Grade", "Tags"],
    ),
    (
        "Categorias dos produtos",
        "Produtos em categorias erradas atrapalham a navegacao e tambem podem passar a impressao de catalogo desorganizado.",
        ["Categoria", "Categorias"],
    ),
    (
        "Imagem, URL e slug",
        "Problemas aqui confundem o cliente e tambem podem atrapalhar SEO, links partilhados e leitura do Google.",
        ["Imagem", "URL/slug"],
    ),
    (
        "Organizacao interna",
        "Nao e tudo erro grave, mas sao pontos que deixam a loja mais dificil de manter e corrigir no futuro.",
        ["SKU", "Duplicado", "Promocao antiga", "Padronizacao"],
    ),
]

DESCRIPTION_GROUPS = [
    (
        "Restos de IA e texto colado",
        "Casos em que a descricao tem sinais claros de texto gerado/colado sem revisao.",
        ["Descricao/IA"],
    ),
    (
        "Campos com confirmar fornecedor",
        "Descricoes que deixam aparecer texto interno como RAM: confirmar fornecedor, Ecrã: confirmar fornecedor ou GTIN/EAN: confirmar fornecedor.",
        ["Descricao/confirmar fornecedor"],
    ),
    (
        "Avaliacoes genericas ou possivelmente falsas",
        "Descricoes que falam em avaliacao/4.8, mas o produto aparece sem reviews na API.",
        ["Descricao/avaliacao"],
    ),
    (
        "Campos vazios na descricao",
        "Textos com campos como Bateria ou Versao aparecendo sem informacao preenchida.",
        ["Descricao/campos vazios"],
    ),
    (
        "Imagem externa dentro da descricao",
        "Descricoes que usam imagem de outro site em vez de imagem propria.",
        ["Descricao/imagem externa"],
    ),
    (
        "Especificacao ou grade contraditoria",
        "Pontos em que a descricao parece nao bater com o titulo ou com a grade do produto.",
        ["Descricao/especificacao", "Descricao/Grade"],
    ),
    (
        "Descricao ausente ou template ruim",
        "Produtos sem descricao publica ou com texto de template mal aplicado.",
        ["Descricao"],
    ),
]

TYPE_INFO = {
    "Marca": ("Alta", "Corrigir marca/filtro/icone."),
    "Tag/condicao": ("Alta", "Corrigir tag de Novo, Usado ou Open Box."),
    "Tag/Grade": ("Alta", "Usar a tag de grade correta, por exemplo Usado - Grade A."),
    "Tags": ("Media", "Adicionar uma tag de condicao."),
    "Categoria": ("Alta", "Mover o produto para a categoria correta."),
    "Categorias": ("Media", "Adicionar categoria correta."),
    "Imagem": ("Alta", "Corrigir imagem, alt text ou nome do ficheiro."),
    "URL/slug": ("Media", "Corrigir o slug/URL."),
    "SKU": ("Baixa", "Adicionar codigo interno/referencia do produto. Se existir EAN/GTIN, preencher tambem."),
    "Duplicado": ("Media", "Unificar ou diferenciar estoque, unidade ou localizacao."),
    "Promocao antiga": ("Media", "Remover de campanha antiga ou atualizar campanha."),
    "Padronizacao": ("Baixa", "Padronizar idioma do titulo se a loja quiser manter tudo em PT."),
    "Descricao/IA": ("Alta", "Limpar o resto de IA e revisar o texto."),
    "Descricao/confirmar fornecedor": ("Alta", "Trocar por informacao real ou remover estes campos antes de publicar."),
    "Descricao/avaliacao": ("Alta", "Remover avaliacao generica ou usar reviews reais."),
    "Descricao/campos vazios": ("Media", "Preencher ou remover campos vazios."),
    "Descricao/imagem externa": ("Media", "Hospedar imagem propria ou remover."),
    "Descricao/especificacao": ("Alta", "Corrigir a especificacao tecnica."),
    "Descricao/Grade": ("Alta", "Unificar grade no titulo e descricao."),
    "Descricao": ("Media", "Criar ou revisar a descricao do produto."),
}

KIND_DISPLAY = {
    "SKU": "Codigo interno / SKU",
}


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(32, 32, 32)

    doc.styles["Heading 1"].font.size = Pt(22)
    doc.styles["Heading 2"].font.size = Pt(21)
    doc.styles["Heading 3"].font.size = Pt(18)
    doc.styles["Heading 4"].font.size = Pt(14)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def set_size(paragraph, size: float):
    for run in paragraph.runs:
        run.font.size = Pt(size)


def add_big_heading(doc: Document, text: str, level: int = 2, size: float = 22, center: bool = False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(20 if level <= 2 else 16)
    p.paragraph_format.space_after = Pt(9 if level <= 2 else 7)
    p.paragraph_format.keep_with_next = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 105, 0)
    return p


def add_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    set_size(p, 9.5)


def add_labeled_line(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run(label).bold = True
    p.add_run(text)
    set_size(p, 9.2)
    return p


def display_kind(kind: str) -> str:
    return KIND_DISPLAY.get(kind, kind)


def display_detail(item: dict) -> str:
    if item["kind"] == "SKU":
        return "Produto sem codigo interno/referencia nos dados do site."
    return item["detail"]


def display_expected(item: dict) -> str:
    if item["kind"] == "SKU":
        return "Adicionar codigo interno do produto. Se existir EAN/GTIN, preencher tambem."
    return item.get("expected", "")


def add_link_line(doc: Document, url: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run("Link: ").bold = True
    add_hyperlink(p, url, url)
    set_size(p, 8.8)
    return p


def split_issues(products: list[dict], issues: dict[int, list[dict]]):
    products_by_id = {product["id"]: product for product in products}
    boss_by_kind: dict[str, list[tuple[dict, dict]]] = defaultdict(list)
    desc_by_kind: dict[str, list[tuple[dict, dict]]] = defaultdict(list)

    for product_id, items in issues.items():
        product = products_by_id[product_id]
        for item in items:
            target = desc_by_kind if item["kind"] in DESCRIPTION_KINDS else boss_by_kind
            target[item["kind"]].append((product, item))

    for mapping in [boss_by_kind, desc_by_kind]:
        for kind in mapping:
            mapping[kind].sort(key=lambda row: (clean(row[0].get("name")).lower(), row[0]["id"]))

    return boss_by_kind, desc_by_kind


def remove_deductive_issues(issues: dict[int, list[dict]]) -> dict[int, list[dict]]:
    FILTERED_DEDUCTIONS.clear()
    filtered: dict[int, list[dict]] = {}
    for product_id, items in issues.items():
        kept = []
        for item in items:
            detail = item.get("detail", "")
            if item["kind"] in DEDUCTIVE_KINDS:
                FILTERED_DEDUCTIONS.append((product_id, item["kind"], "removido porque dependia de deducao ou regra interna de tag/estado/grade"))
                continue
            if item["kind"] == "Marca" and "a confirmar" in detail.lower():
                FILTERED_DEDUCTIONS.append((product_id, item["kind"], "removido porque era caso de marca/filtro apenas a confirmar"))
                continue
            kept.append(item)
        if kept:
            filtered[product_id] = kept
    return filtered


def add_supplier_placeholder_issues(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, list[dict]]:
    enriched = {product_id: list(items) for product_id, items in issues.items()}
    pattern = re.compile(r"[^\n.]*confirmar\s+(?:o\s+)?(?:fornecedor|vendedor)[^\n.]*", re.I)

    for product in products:
        raw_desc = (product.get("short_description") or "") + "\n" + (product.get("description") or "")
        plain_desc = strip_html(raw_desc)
        matches = []
        for match in pattern.findall(plain_desc):
            text = clean(match).strip(" -•\t")
            text = re.sub(r"\s+", " ", text)
            if text and text not in matches:
                matches.append(text[:140])

        if not matches:
            continue

        examples = "; ".join(matches[:6])
        extra = len(matches) - 6
        if extra > 0:
            examples += f"; e mais {extra} campo(s)"

        item = {
            "kind": "Descricao/confirmar fornecedor",
            "detail": f"Descricao mostra placeholder interno publicado: {examples}",
            "expected": "Preencher com informacao real ou remover esses campos. Cliente nao deve ver 'confirmar fornecedor/vendedor'.",
        }
        if item not in enriched[product["id"]]:
            enriched.setdefault(product["id"], []).append(item)

    return enriched


def count_rows(mapping: dict[str, list[tuple[dict, dict]]], kinds: list[str] | None = None) -> int:
    if kinds is None:
        return sum(len(rows) for rows in mapping.values())
    return sum(len(mapping.get(kind, [])) for kind in kinds)


def count_products(mapping: dict[str, list[tuple[dict, dict]]]) -> int:
    return len({product["id"] for rows in mapping.values() for product, _ in rows})


def add_summary(doc: Document, products: list[dict], issues, boss_by_kind, desc_by_kind, empty_categories, security):
    add_big_heading(doc, "Resumo rapido", level=2, size=20)
    lines = [
        f"Produtos analisados: {len(products)}",
        f"Produtos com pelo menos um problema: {len(issues)}",
        f"Problemas totais de produto: {sum(len(items) for items in issues.values())}",
        f"Problemas fora de descricao, para a chefia: {count_rows(boss_by_kind)}",
        f"Produtos com problema fora de descricao: {count_products(boss_by_kind)}",
        f"Problemas de descricao, para o colega revisar: {count_rows(desc_by_kind)}",
        f"Produtos com descricao para revisar: {count_products(desc_by_kind)}",
        f"Categorias vazias confirmadas: {len(empty_categories)}",
        f"Pontos de seguranca/configuracao confirmados: {len(security)}",
        f"Itens removidos na revisao para evitar falso positivo: {len(REVIEW_REMOVED)}",
        f"Itens removidos por dependerem de deducao: {len(FILTERED_DEDUCTIONS)}",
        "Imagens no documento: 0",
        "Tabelas no documento: 0",
    ]
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")

    add_big_heading(doc, "Como usar este documento", level=2, size=20)
    add_note(
        doc,
        "A primeira parte foi pensada para a chefia: mostra o que afeta catalogo, filtros, organizacao, menu e seguranca.",
    )
    add_note(
        doc,
        "A segunda parte e separada so para descricoes de produtos. Assim quem for corrigir texto nao precisa passar por seguranca, tags, marcas ou categorias.",
    )
    add_note(
        doc,
        "Alguns produtos aparecem nas duas partes quando tem, por exemplo, erro de marca e tambem erro de descricao. Em cada parte aparece apenas o problema daquela area.",
    )
    add_note(
        doc,
        "Nesta versao eu removi problemas baseados em deducao de produto novo/usado, grade, falta de tag de condicao ou regra subjetiva de estilo. Ficaram apenas pontos mais objetivos.",
    )


def add_issue_block(doc: Document, product: dict, item: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{clean(product.get('name'))} | ID {product['id']}").bold = True
    set_size(p, 10)

    add_link_line(doc, product["permalink"])
    add_labeled_line(doc, "Problema: ", display_detail(item))
    expected = display_expected(item)
    if expected:
        add_labeled_line(doc, "Corrigir: ", expected)


def add_kind_section(doc: Document, kind: str, rows: list[tuple[dict, dict]]):
    priority, fix = TYPE_INFO.get(kind, ("Media", "Revisar manualmente."))
    add_big_heading(doc, f"{display_kind(kind)} ({len(rows)})", level=4, size=15)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Prioridade: ").bold = True
    p.add_run(priority)
    p.add_run(" | Acao sugerida: ").bold = True
    p.add_run(fix)
    set_size(p, 9.2)

    for product, item in rows:
        add_issue_block(doc, product, item)


def add_grouped_section(doc: Document, title: str, intro: str, groups, mapping):
    add_big_heading(doc, title, level=2, size=23, center=True)
    add_note(doc, intro)

    for group_title, group_desc, kinds in groups:
        total = count_rows(mapping, kinds)
        if total == 0:
            continue
        doc.add_page_break()
        add_big_heading(doc, f"Inicio da categoria: {group_title} ({total})", level=3, size=21, center=True)
        add_note(doc, group_desc)
        for kind in kinds:
            rows = mapping.get(kind, [])
            if rows:
                add_kind_section(doc, kind, rows)


def add_general_for_boss(doc: Document, empty_categories, security):
    doc.add_page_break()
    add_big_heading(doc, "Inicio da categoria: Menu e categorias vazias", level=3, size=21, center=True)
    add_note(doc, "Estas categorias aparecem no menu, mas abrem paginas sem produtos. Para o cliente, isso parece area abandonada.")
    for name, url in empty_categories:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run(name).bold = True
        set_size(p, 10)
        add_link_line(doc, url)
        add_labeled_line(doc, "Problema: ", "Categoria aparece no menu e esta vazia/sem produtos.")
        add_labeled_line(doc, "Corrigir: ", "Remover do menu ou esconder ate ter produtos.")

    doc.add_page_break()
    add_big_heading(doc, "Inicio da categoria: Seguranca e configuracao", level=3, size=21, center=True)
    add_note(doc, "Estes pontos foram vistos por consulta publica/passiva. Nao e um teste invasivo, mas ja mostra coisas que valem corrigir.")
    for title, detail, expected, url in security:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run(title).bold = True
        set_size(p, 10)
        add_link_line(doc, url)
        add_labeled_line(doc, "Problema: ", detail)
        add_labeled_line(doc, "Corrigir: ", expected)


def build_doc(products, issues, empty_categories, security):
    boss_by_kind, desc_by_kind = split_issues(products, issues)

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Clickmed.pt - relatorio separado por responsabilidade", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_size(title, 24)
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Versao humanizada, sem imagens e sem tabelas.").italic = True

    add_summary(doc, products, issues, boss_by_kind, desc_by_kind, empty_categories, security)

    doc.add_page_break()
    add_grouped_section(
        doc,
        "Parte 1 - Para a chefia",
        "Aqui ficam os problemas de loja/catalogo que normalmente dependem de decisao, prioridade ou correcao de administracao.",
        BOSS_GROUPS,
        boss_by_kind,
    )
    add_general_for_boss(doc, empty_categories, security)

    doc.add_page_break()
    add_grouped_section(
        doc,
        "Parte 2 - Apenas descricoes dos produtos",
        "Esta parte e so para quem vai revisar texto. Nao coloquei aqui marca, tag, categoria, seguranca ou organizacao interna.",
        DESCRIPTION_GROUPS,
        desc_by_kind,
    )

    doc.save(OUT_DOCX)
    shutil.copyfile(OUT_DOCX, LATEST_DOCX)
    shutil.copyfile(OUT_DOCX, FINAL_BLOCK_DOCX)


def md_issue_block(lines: list[str], product: dict, item: dict):
    lines.append(f"#### {clean(product.get('name'))} | ID {product['id']}")
    lines.append(f"Link: {product['permalink']}")
    lines.append(f"Problema: {display_detail(item)}")
    expected = display_expected(item)
    if expected:
        lines.append(f"Corrigir: {expected}")
    lines.append("")


def build_md(products, issues, empty_categories, security):
    boss_by_kind, desc_by_kind = split_issues(products, issues)
    lines = [
        "# Clickmed.pt - relatorio separado por responsabilidade",
        "",
        "Versao humanizada, sem imagens e sem tabelas.",
        "",
        "## Resumo rapido",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com pelo menos um problema: {len(issues)}",
        f"- Problemas totais de produto: {sum(len(items) for items in issues.values())}",
        f"- Problemas fora de descricao, para a chefia: {count_rows(boss_by_kind)}",
        f"- Produtos com problema fora de descricao: {count_products(boss_by_kind)}",
        f"- Problemas de descricao, para o colega revisar: {count_rows(desc_by_kind)}",
        f"- Produtos com descricao para revisar: {count_products(desc_by_kind)}",
        f"- Categorias vazias confirmadas: {len(empty_categories)}",
        f"- Pontos de seguranca/configuracao confirmados: {len(security)}",
        f"- Itens removidos na revisao para evitar falso positivo: {len(REVIEW_REMOVED)}",
        f"- Itens removidos por dependerem de deducao: {len(FILTERED_DEDUCTIONS)}",
        f"- Imagens no documento: 0",
        f"- Tabelas no documento: 0",
        "",
        "## PARTE 1 - PARA A CHEFIA",
        "",
    ]

    for group_title, group_desc, kinds in BOSS_GROUPS:
        total = count_rows(boss_by_kind, kinds)
        if total == 0:
            continue
        lines += [f"### INICIO DA CATEGORIA: {group_title.upper()} ({total})", "", group_desc, ""]
        for kind in kinds:
            rows = boss_by_kind.get(kind, [])
            if not rows:
                continue
            priority, fix = TYPE_INFO[kind]
            lines += [f"#### {display_kind(kind).upper()} ({len(rows)})", f"Prioridade: {priority}", f"Acao sugerida: {fix}", ""]
            for product, item in rows:
                md_issue_block(lines, product, item)

    lines += ["### INICIO DA CATEGORIA: MENU E CATEGORIAS VAZIAS", ""]
    for name, url in empty_categories:
        lines += [f"#### {name}", f"Link: {url}", "Problema: Categoria aparece no menu e esta vazia/sem produtos.", "Corrigir: Remover do menu ou esconder ate ter produtos.", ""]

    lines += ["### INICIO DA CATEGORIA: SEGURANCA E CONFIGURACAO", "", "Pontos confirmados por consulta publica/passiva. Nao e teste invasivo.", ""]
    for title, detail, expected, url in security:
        lines += [f"#### {title}", f"Link: {url}", f"Problema: {detail}", f"Corrigir: {expected}", ""]

    lines += [
        "## PARTE 2 - APENAS DESCRICOES DOS PRODUTOS",
        "",
        "Esta parte e so para quem vai revisar texto. Nao coloquei aqui marca, tag, categoria, seguranca ou organizacao interna.",
        "",
    ]
    for group_title, group_desc, kinds in DESCRIPTION_GROUPS:
        total = count_rows(desc_by_kind, kinds)
        if total == 0:
            continue
        lines += [f"### INICIO DA CATEGORIA: {group_title.upper()} ({total})", "", group_desc, ""]
        for kind in kinds:
            rows = desc_by_kind.get(kind, [])
            if not rows:
                continue
            priority, fix = TYPE_INFO[kind]
            lines += [f"#### {display_kind(kind).upper()} ({len(rows)})", f"Prioridade: {priority}", f"Acao sugerida: {fix}", ""]
            for product, item in rows:
                md_issue_block(lines, product, item)

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    shutil.copyfile(OUT_MD, LATEST_MD)
    shutil.copyfile(OUT_MD, FINAL_BLOCK_MD)


def validate(products, issues, empty_categories, security):
    doc = Document(OUT_DOCX)
    boss_by_kind, desc_by_kind = split_issues(products, issues)
    with zipfile.ZipFile(OUT_DOCX) as archive:
        bad = archive.testzip()

    lines = [
        "# Validacao do relatorio humanizado separado",
        "",
        f"- DOCX integro: {'sim' if bad is None else 'nao'}",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com problemas: {len(issues)}",
        f"- Problemas totais de produto: {sum(len(items) for items in issues.values())}",
        f"- Problemas fora de descricao: {count_rows(boss_by_kind)}",
        f"- Produtos com problema fora de descricao: {count_products(boss_by_kind)}",
        f"- Problemas de descricao: {count_rows(desc_by_kind)}",
        f"- Produtos com descricao para revisar: {count_products(desc_by_kind)}",
        f"- Categorias vazias confirmadas: {len(empty_categories)}",
        f"- Seguranca/configuracao confirmados: {len(security)}",
        f"- Itens removidos por dependerem de deducao: {len(FILTERED_DEDUCTIONS)}",
        f"- Tabelas no DOCX: {len(doc.tables)}",
        f"- Imagens no DOCX: {len(doc.inline_shapes)}",
        f"- Paragrafos no DOCX: {len(doc.paragraphs)}",
        "- Arquivo atualizado tambem em relatorio_clickmed_FINAL_REVISADO.docx: sim",
        "- Arquivo atualizado tambem em relatorio_clickmed_final_blocos.docx: sim",
    ]
    OUT_VALIDATION.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_review_products()
    issues = add_supplier_placeholder_issues(products, remove_deductive_issues(reviewed_issues(products)))
    empty_categories = confirmed_empty_categories()
    security = confirmed_security_findings()
    build_doc(products, issues, empty_categories, security)
    build_md(products, issues, empty_categories, security)
    validate(products, issues, empty_categories, security)
    print(OUT_DOCX)
    print(OUT_MD)
    print(OUT_VALIDATION)


if __name__ == "__main__":
    main()
