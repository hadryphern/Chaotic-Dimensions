from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gerar_relatorio_completo_docx import (
    EMPTY_MENU_CATEGORIES,
    SECURITY_FINDINGS,
    add_hyperlink,
    clean,
    collect_product_issues,
    load_products,
    make_non_product_card,
    make_product_card,
    product_brands,
    product_categories,
    product_tags,
)
from gerar_relatorio_humanizado_docx import GROUPS, GROUP_TEXT, grouped_rows


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_blocos_com_imagens.docx"
OUT_MD = ROOT / "relatorio_clickmed_blocos_com_imagens.md"


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.color.rgb = RGBColor(35, 35, 35)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)


def add_divider(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("")
    run.add_break()
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.space_before = Pt(8)


def add_problem_lines(doc: Document, items: list[dict]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{item['kind']}: ").bold = True
        p.add_run(item["detail"])
        if item.get("expected"):
            p.add_run(" | Corrigir: ").bold = True
            p.add_run(item["expected"])


def add_product_block(doc: Document, product: dict, items: list[dict], card_path: Path):
    title = doc.add_heading(clean(product.get("name")), level=3)
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(4)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    meta.add_run("Problema(s): ").bold = True
    meta.add_run(f"{len(items)} encontrado(s)")

    add_problem_lines(doc, items)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Link: ").bold = True
    add_hyperlink(p, product["permalink"], product["permalink"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Tags: ").bold = True
    p.add_run(", ".join(product_tags(product)) or "sem tag")
    p.add_run(" | Marcas: ").bold = True
    p.add_run(", ".join(product_brands(product)) or "sem marca")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("Categorias: ").bold = True
    p.add_run(", ".join(product_categories(product)) or "sem categoria")

    doc.add_picture(str(card_path), width=Inches(6.55))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(20)

    add_divider(doc)


def make_cards(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, Path]:
    products_by_id = {product["id"]: product for product in products}
    cards = {}
    for product_id in sorted(issues):
        cards[product_id] = make_product_card(products_by_id[product_id], issues[product_id])
    return cards


def make_general_cards() -> list[tuple[Path, str, str, str]]:
    cards = []
    index = 1
    for name, url in EMPTY_MENU_CATEGORIES:
        detail = "Categoria aparece no menu, mas esta vazia/sem produtos."
        card = make_non_product_card(name, url, detail, "Remover do menu ou esconder ate ter produtos.", index)
        cards.append((card, name, detail, url))
        index += 1

    for kind, detail, expected in SECURITY_FINDINGS:
        title = f"Seguranca - {kind}"
        card = make_non_product_card(title, "https://clickmed.pt/", detail, expected, index)
        cards.append((card, title, f"{detail} Corrigir: {expected}", "https://clickmed.pt/"))
        index += 1
    return cards


def build_doc(products: list[dict], issues: dict[int, list[dict]]):
    grouped = grouped_rows(products, issues)
    cards = make_cards(products, issues)
    products_by_id = {product["id"]: product for product in products}

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Clickmed.pt - problemas encontrados", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Formato em blocos: nome, problema, link e imagem juntos. Sem tabelas.").italic = True

    doc.add_heading("Resumo", level=2)
    for text in [
        f"Produtos analisados: {len(products)}",
        f"Produtos com erro/inconveniente: {len(issues)}",
        f"Total de problemas encontrados: {sum(len(items) for items in issues.values())}",
        f"Categorias vazias no menu: {len(EMPTY_MENU_CATEGORIES)}",
        f"Pontos de seguranca/configuracao: {len(SECURITY_FINDINGS)}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph(
        "Cada produto abaixo tem o nome como titulo, os problemas logo em seguida, o link completo escrito por extenso e a imagem-evidencia junto do mesmo bloco."
    )

    doc.add_page_break()

    for group in GROUPS:
        rows = grouped.get(group, [])
        doc.add_heading(group, level=2)
        doc.add_paragraph(GROUP_TEXT[group]["intro"])
        doc.add_paragraph(f"Como corrigir: {GROUP_TEXT[group]['fix']}")
        doc.add_paragraph(f"Produtos nesta parte: {len(rows)}")

        for product, items in rows:
            add_product_block(doc, product, items, cards[product["id"]])

        doc.add_page_break()

    doc.add_heading("Categorias vazias no menu e seguranca", level=2)
    for card, title_text, detail, link in make_general_cards():
        heading = doc.add_heading(title_text, level=3)
        heading.paragraph_format.space_before = Pt(14)

        p = doc.add_paragraph()
        p.add_run("Problema: ").bold = True
        p.add_run(detail)

        p = doc.add_paragraph()
        p.add_run("Link: ").bold = True
        add_hyperlink(p, link, link)

        doc.add_picture(str(card), width=Inches(6.55))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(20)
        add_divider(doc)

    doc.save(OUT_DOCX)


def build_md(products: list[dict], issues: dict[int, list[dict]]):
    grouped = grouped_rows(products, issues)
    lines = [
        "# Clickmed.pt - problemas encontrados",
        "",
        "Formato em blocos: nome, problema, link e imagem juntos. Sem tabelas.",
        "",
        "## Resumo",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com erro/inconveniente: {len(issues)}",
        f"- Total de problemas encontrados: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias no menu: {len(EMPTY_MENU_CATEGORIES)}",
        f"- Pontos de seguranca/configuracao: {len(SECURITY_FINDINGS)}",
        "",
    ]

    for group in GROUPS:
        rows = grouped.get(group, [])
        lines += [f"## {group}", "", GROUP_TEXT[group]["intro"], "", f"Produtos nesta parte: {len(rows)}", ""]
        for product, items in rows:
            lines.append(f"### {clean(product.get('name'))}")
            lines.append("")
            lines.append(f"Link: {product['permalink']}")
            lines.append("")
            lines.append(f"Tags: {', '.join(product_tags(product)) or 'sem tag'}")
            lines.append(f"Marcas: {', '.join(product_brands(product)) or 'sem marca'}")
            lines.append(f"Categorias: {', '.join(product_categories(product)) or 'sem categoria'}")
            lines.append("")
            lines.append("Problemas:")
            for item in items:
                line = f"- {item['kind']}: {item['detail']}"
                if item.get("expected"):
                    line += f" | Corrigir: {item['expected']}"
                lines.append(line)
            lines.append("")
            lines.append("---")
            lines.append("")

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_products()
    issues = collect_product_issues(products)
    build_doc(products, issues)
    build_md(products, issues)
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
