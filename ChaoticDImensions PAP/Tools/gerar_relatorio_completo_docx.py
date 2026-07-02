from __future__ import annotations

import csv
import glob
import html
import json
import re
import time
import unicodedata
import urllib.request
from collections import Counter, defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_COMPLETO_todos_erros.docx"
OUT_CSV = ROOT / "clickmed_erros_completos.csv"
CARD_DIR = ROOT / "evidencias_clickmed_todos_produtos"
THUMB_DIR = CARD_DIR / "_thumbs"

PRODUCT_API = "https://clickmed.pt/wp-json/wc/store/v1/products?per_page=100&page={page}"


BRAND_ALIASES = {
    "Apple": [r"\bapple\b", r"\biphone\b", r"\bipad\b", r"\bmacbook\b", r"\bairpods\b", r"\bimac\b", r"\bapple watch\b"],
    "Samsung": [r"\bsamsung\b", r"\bgalaxy\b"],
    "Xiaomi": [r"\bxiaomi\b", r"\bredmi\b", r"\bpoco\b"],
    "Sony": [r"\bsony\b", r"\bplaystation\b", r"\bps5\b", r"\bps4\b", r"\bdualsense\b"],
    "Nintendo": [r"\bnintendo\b", r"\bswitch\b"],
    "HP": [r"\bhp\b"],
    "Acer": [r"\bacer\b"],
    "Lenovo": [r"\blenovo\b"],
    "Asus": [r"\basus\b"],
    "LG": [r"\blg\b"],
    "Oppo": [r"\boppo\b"],
    "Razer": [r"\brazer\b"],
    "Corsair": [r"\bcorsair\b"],
    "JBL": [r"\bjbl\b"],
    "Meta": [r"\bmeta\b", r"\bquest\b"],
    "MSI": [r"\bmsi\b"],
    "Hisense": [r"\bhisense\b"],
    "Haier": [r"\bhaier\b"],
    "Starlink": [r"\bstarlink\b"],
    "Ulefone": [r"\bulefone\b"],
    "Fantech": [r"\bfantech\b"],
    "Nox": [r"\bnox\b"],
    "Alpha Gamer": [r"\balpha gamer\b"],
    "Floe": [r"\bfloe\b"],
    "Honor": [r"\bhonor\b"],
    "Nothing": [r"\bnothing\b"],
    "Toshiba": [r"\btoshiba\b"],
    "Silicon Power": [r"\bsilicon power\b"],
    "Thrustmaster": [r"\bthrustmaster\b"],
    "Logitech": [r"\blogitech\b"],
}
KNOWN_BRANDS = set(BRAND_ALIASES)

ENGLISH_TERMS = [
    "Black",
    "Blue",
    "White",
    "Green",
    "Red",
    "Purple",
    "Silver",
    "Gold",
    "Orange",
    "Grey",
    "Gray",
    "Midnight",
    "Starlight",
    "Desert Titanium",
    "Natural Titanium",
    "Space Gray",
    "Cosmic Orange",
]

OLD_PROMO_CATS = {
    "Promocoes Marco",
    "Promoções Março",
    "Promoções Abril",
    "Promoções Maio",
    "Promoções Junho",
}

EMPTY_MENU_CATEGORIES = [
    ("Tablets Usados", "https://clickmed.pt/categoria-produto/tablets/tablets-usados/"),
    ("Teclados", "https://clickmed.pt/categoria-produto/perifericos/teclados/"),
    ("Gamepad", "https://clickmed.pt/categoria-produto/perifericos/gamepad/"),
    ("Fontes de alimentação", "https://clickmed.pt/categoria-produto/componentes/fontes-de-alimentacao/"),
    ("Motherboard", "https://clickmed.pt/categoria-produto/componentes/motherboard/"),
    ("Processadores", "https://clickmed.pt/categoria-produto/componentes/processadores/"),
    ("Mesa Gaming", "https://clickmed.pt/categoria-produto/gaming/mesa-gaming/"),
    ("Smart Home", "https://clickmed.pt/categoria-produto/smart-home/"),
    ("Robots Aspiradores", "https://clickmed.pt/categoria-produto/smart-home/robot-aspirador/"),
    ("Assistentes Virtuais", "https://clickmed.pt/categoria-produto/smart-home/assistentes-virtuais/"),
    ("Câmaras de Segurança", "https://clickmed.pt/categoria-produto/smart-home/camaras-de-seguranca/"),
    ("Microfones", "https://clickmed.pt/categoria-produto/audio/microfones/"),
    ("Impressoras", "https://clickmed.pt/categoria-produto/impressoras/"),
    ("Espaço Kids", "https://clickmed.pt/categoria-produto/espaco-kids/"),
]

SECURITY_FINDINGS = [
    ("Headers", "Homepage/loja sem HSTS visivel", "Adicionar Strict-Transport-Security em todas as paginas HTTPS."),
    ("Headers", "Homepage/loja sem Content-Security-Policy geral", "Definir CSP e pelo menos frame-ancestors 'self'."),
    ("Headers", "Homepage/loja sem Referrer-Policy visivel", "Usar Referrer-Policy: strict-origin-when-cross-origin."),
    ("Headers", "Homepage/loja sem Permissions-Policy visivel", "Bloquear camera/microfone/geolocalizacao se nao forem usados."),
    ("WordPress", "xmlrpc.php acessivel", "Desativar XML-RPC se nao for usado ou proteger com WAF/rate limit."),
    ("WordPress", "Autor admin exposto no JSON-LD", "Evitar usar/expor usuario chamado admin."),
    ("WordPress", "Versoes de plugins/tema expostas no HTML", "Reduzir exposicao de versoes e manter tudo atualizado."),
    ("API", "REST API publica mostra muitas rotas: WooCommerce, Sequra, Klarna, Ifthenpay, AI Engine, MCP", "Desativar o que nao for usado e garantir auth/rate-limit."),
    ("Pagamentos", "Gateways e webhooks aparecem publicamente por rotas REST", "Confirmar assinatura/secret nos webhooks e nunca confiar em valor vindo do front-end."),
    ("Conta/Admin", "wp-login.php publico", "Ativar 2FA, WAF/rate-limit, senha forte e usuarios com menor privilegio possivel."),
]


def clean(value: str | None) -> str:
    text = html.unescape(value or "").replace("\xa0", " ").strip()
    replacements = {
        "â€“": "-",
        "â€”": "-",
        "â€˜": "'",
        "â€™": "'",
        "â€œ": '"',
        "â€": '"',
        "Ã¡": "á",
        "Ã ": "à",
        "Ã¢": "â",
        "Ã£": "ã",
        "Ã©": "é",
        "Ãª": "ê",
        "Ã­": "í",
        "Ã³": "ó",
        "Ã´": "ô",
        "Ãµ": "õ",
        "Ãº": "ú",
        "Ã§": "ç",
        "Ã": "Á",
        "Ã‰": "É",
        "Ã“": "Ó",
        "Ã‡": "Ç",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def card_text(value: str | None) -> str:
    text = clean(value)
    text = text.replace("–", "-").replace("—", "-").replace("“", '"').replace("”", '"').replace("’", "'")
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")


def strip_html(value: str | None) -> str:
    text = re.sub(r"<br\s*/?>", "\n", value or "", flags=re.I)
    text = re.sub(r"</(p|li|div|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = clean(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s+", "\n", text)
    return text.strip()


def load_products() -> list[dict]:
    files = sorted(glob.glob("/tmp/clickmed_products_page*.json"))
    products: list[dict] = []
    if files:
        for file_name in files:
            products.extend(json.load(open(file_name, encoding="utf-8")))
        return products

    for page in range(1, 6):
        req = urllib.request.Request(PRODUCT_API.format(page=page), headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as response:
            page_products = json.loads(response.read().decode("utf-8"))
        products.extend(page_products)
        time.sleep(0.5)
    return products


def product_tags(product: dict) -> list[str]:
    return [clean(tag.get("name")) for tag in product.get("tags", [])]


def product_categories(product: dict) -> list[str]:
    return [clean(cat.get("name")) for cat in product.get("categories", [])]


def product_brands(product: dict) -> list[str]:
    return [clean(brand.get("name")) for brand in product.get("brands", [])]


def detect_brand(text: str) -> set[str]:
    lower = clean(text).lower()
    found = set()
    for brand, patterns in BRAND_ALIASES.items():
        if any(re.search(pattern, lower) for pattern in patterns):
            found.add(brand)
    return found


def detect_grade(text: str) -> str | None:
    match = re.search(r"grade\s*[-–]?\s*([abc])", clean(text), re.I)
    return match.group(1).upper() if match else None


def detect_condition(text: str) -> str | None:
    lower = clean(text).lower()
    if "open box" in lower:
        return "OPEN BOX"
    if re.search(r"\busad[oa]\b|\brecondicionado\b", lower):
        return "Usado"
    if re.search(r"\bnov[oa]\b", lower):
        return "Novo"
    return None


def memory_tokens(text: str) -> set[str]:
    lower = clean(text).lower()
    return {
        f"{match.group(1)}{match.group(2)}"
        for match in re.finditer(r"(?<!\d)(1|2|3|4|6|8|12|16|24|32|64|128|256|512)\s*[-/]?\s*(gb|tb)", lower)
    }


def add_issue(issue_map: dict[int, list[dict]], product: dict, kind: str, detail: str, expected: str = ""):
    product_id = product["id"]
    item = {"kind": kind, "detail": detail, "expected": expected}
    if item not in issue_map[product_id]:
        issue_map[product_id].append(item)


def collect_product_issues(products: list[dict]) -> dict[int, list[dict]]:
    issues: dict[int, list[dict]] = defaultdict(list)
    name_counter = Counter(clean(product.get("name")).lower() for product in products)
    duplicate_links: dict[str, list[str]] = defaultdict(list)
    for product in products:
        duplicate_links[clean(product.get("name")).lower()].append(product["permalink"])

    for product in products:
        name = clean(product.get("name"))
        slug = clean(product.get("slug"))
        raw_desc = (product.get("short_description") or "") + "\n" + (product.get("description") or "")
        plain_desc = strip_html(raw_desc)
        all_text = f"{name}\n{plain_desc}"
        tags = product_tags(product)
        cats = product_categories(product)
        brands = product_brands(product)

        title_grade = detect_grade(name)
        if title_grade and not any(title_grade in tag for tag in tags):
            add_issue(
                issues,
                product,
                "Tag/Grade",
                f"Titulo indica Grade {title_grade}, mas tags atuais sao: {', '.join(tags) or 'sem tag'}",
                f"Usar tag Usado - Grade {title_grade}",
            )

        title_condition = detect_condition(name)
        tag_condition = detect_condition(" ".join(tags))
        if title_condition and tag_condition and title_condition != tag_condition:
            add_issue(
                issues,
                product,
                "Tag/condicao",
                f"Titulo indica {title_condition}, mas tag indica {tag_condition}",
                "Corrigir tag de condicao do produto",
            )

        cat_new = any(re.search(r"\bNovos?\b|\bNovo\b", cat, re.I) and cat != "Novidades" for cat in cats)
        cat_used = any(re.search(r"\bUsados?\b|\bUsado\b", cat, re.I) for cat in cats)
        product_used = title_condition == "Usado" or tag_condition == "Usado" or bool(title_grade)
        product_new = title_condition == "Novo" or tag_condition == "Novo"
        if cat_new and product_used:
            add_issue(issues, product, "Categoria", f"Produto parece usado, mas esta em categoria de novo: {', '.join(cats)}", "Mover para categoria usada correta")
        if cat_used and product_new:
            add_issue(issues, product, "Categoria", f"Produto parece novo, mas esta em categoria de usados: {', '.join(cats)}", "Mover para categoria nova correta")

        desc_grade = detect_grade(plain_desc)
        if title_grade and desc_grade and title_grade != desc_grade:
            add_issue(issues, product, "Descricao/Grade", f"Titulo diz Grade {title_grade}, mas descricao diz Grade {desc_grade}", "Unificar grade no titulo e descricao")

        title_brands = detect_brand(name)
        assigned_known = set(brands) & KNOWN_BRANDS
        if title_brands and assigned_known and not (title_brands & assigned_known):
            add_issue(issues, product, "Marca", f"Marca do titulo {', '.join(sorted(title_brands))}; marca atribuida {', '.join(brands)}", "Corrigir marca/filtro/icone")
        elif title_brands and not brands:
            add_issue(issues, product, "Marca", f"Produto tem marca no titulo ({', '.join(sorted(title_brands))}), mas nao tem marca atribuida", "Adicionar marca correta")
        elif title_brands and brands == ["Compatível"]:
            add_issue(issues, product, "Marca", f"Produto tem marca no titulo ({', '.join(sorted(title_brands))}), mas marca atribuida e apenas Compativel", "Separar marca real de atributo compativel")

        name_mem = memory_tokens(name)
        slug_mem = memory_tokens(slug)
        if slug_mem and name_mem and (slug_mem - name_mem):
            add_issue(issues, product, "URL/slug", f"URL tem capacidade {', '.join(sorted(slug_mem - name_mem))}, mas titulo tem {', '.join(sorted(name_mem))}", "Corrigir slug/URL")
        if "agb" in slug.lower():
            add_issue(issues, product, "URL/slug", "Slug contem typo de capacidade: agb", "Corrigir typo no slug")

        for image in product.get("images", []):
            image_text = " ".join([clean(image.get("src")), clean(image.get("name")), clean(image.get("alt"))])
            image_mem = memory_tokens(image_text)
            extra = image_mem - name_mem
            if image_mem and name_mem and extra:
                add_issue(
                    issues,
                    product,
                    "Imagem",
                    f"Imagem/alt indica {', '.join(sorted(image_mem))}, mas titulo indica {', '.join(sorted(name_mem))}",
                    "Corrigir imagem, alt e nome do ficheiro",
                )
                break

        if (product.get("review_count") == 0 and clean(product.get("average_rating")) in {"", "0"}) and (
            "4.8 / 5" in plain_desc or "Baseado em avaliações de clientes" in plain_desc or "Baseado em avaliacoes de clientes" in plain_desc
        ):
            add_issue(issues, product, "Descricao/avaliacao", "Descricao mostra 4.8/5 ou avaliacoes, mas API indica 0 reviews e rating 0", "Remover avaliacao falsa/generica")

        if not plain_desc:
            add_issue(issues, product, "Descricao", "Produto sem descricao publica", "Criar descricao correta")
        if "meses meses" in plain_desc:
            add_issue(issues, product, "Descricao", "Texto duplicado: meses meses", "Corrigir garantia")
        if ":contentReference" in raw_desc:
            add_issue(issues, product, "Descricao/IA", "Descricao contem resto de IA: :contentReference", "Remover lixo de IA e revisar texto")
        if "utm_source=chatgpt" in raw_desc:
            add_issue(issues, product, "Descricao/IA", "Descricao contem link com utm_source=chatgpt", "Reescrever descricao e remover fontes coladas")
        if "www.worten.pt/assetsV4" in raw_desc:
            add_issue(issues, product, "Descricao/imagem externa", "Descricao usa imagem externa da Worten", "Hospedar/corrigir imagem propria ou remover")
        if "A14 Bionic" in plain_desc and not re.search(r"\biPhone\s*12\b|\biPad Air\b", name, re.I):
            add_issue(issues, product, "Descricao/especificacao", "Descricao menciona A14 Bionic em produto que nao parece ser iPhone 12/iPad Air", "Corrigir especificacao tecnica")
        desc_lines = [line.strip(" -•\t") for line in plain_desc.splitlines()]
        if any(line in {"Bateria", "🔋 Bateria", "Versão", "📶 Versão"} for line in desc_lines):
            add_issue(issues, product, "Descricao/campos vazios", "Descricao tem campos vazios como Bateria ou Versao", "Preencher ou remover campos vazios")
        if "open box open box usado" in plain_desc.lower():
            add_issue(issues, product, "Descricao", "Texto repete condicao: Open Box Open Box Usado", "Corrigir template da descricao")

        if not tags:
            add_issue(issues, product, "Tags", "Produto sem tag de condicao", "Adicionar Novo / Open Box / Usado Grade A-B-C")
        if not cats:
            add_issue(issues, product, "Categorias", "Produto sem categoria", "Adicionar categoria correta")
        if not clean(product.get("sku")):
            add_issue(issues, product, "SKU", "Produto sem SKU", "Adicionar SKU/EAN/referencia interna")

        old_promos = [cat for cat in cats if cat in OLD_PROMO_CATS]
        if old_promos:
            add_issue(issues, product, "Promocao antiga", f"Produto ainda em categoria antiga: {', '.join(old_promos)}", "Remover de promocoes antigas ou atualizar campanha")

        english_hits = [term for term in ENGLISH_TERMS if re.search(r"\b" + re.escape(term) + r"\b", name)]
        if english_hits:
            add_issue(issues, product, "Padronizacao", f"Titulo mistura termos/cores em ingles: {', '.join(english_hits)}", "Padronizar idioma do titulo")

        lowered = clean(product.get("name")).lower()
        if name_counter[lowered] > 1:
            others = [link for link in duplicate_links[lowered] if link != product["permalink"]]
            add_issue(issues, product, "Duplicado", f"Nome duplicado com {len(others)} outro(s) produto(s): {' | '.join(others[:3])}", "Unificar ou diferenciar estoque/unidade")

    return issues


def font(size=28, bold=False):
    paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf",
    ]
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = card_text(str(text)).split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=fnt)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def download_thumbnail(product: dict) -> Path | None:
    images = product.get("images") or []
    if not images:
        return None
    url = images[0].get("thumbnail") or images[0].get("src")
    if not url:
        return None
    THUMB_DIR.mkdir(parents=True, exist_ok=True)
    suffix = ".jpg" if ".jpg" in url.lower() or ".jpeg" in url.lower() else ".png"
    path = THUMB_DIR / f"{product['id']}{suffix}"
    if path.exists():
        return path
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as response:
            path.write_bytes(response.read())
        return path
    except Exception:
        return None


def make_product_card(product: dict, issue_items: list[dict]) -> Path:
    CARD_DIR.mkdir(parents=True, exist_ok=True)
    out = CARD_DIR / f"produto_{product['id']}.jpg"
    if out.exists():
        return out

    width = 1200
    height = 430 + min(len(issue_items), 14) * 82
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    header_font = font(36, True)
    name_font = font(31, True)
    body_font = font(28)
    small_font = font(24)
    mono_font = font(22)

    draw.rectangle((0, 0, width, 82), fill=(255, 105, 0))
    draw.text((28, 22), f"Produto com erro - ID {product['id']}", fill="white", font=header_font)

    thumb = download_thumbnail(product)
    thumb_box = (32, 112, 222, 302)
    if thumb and thumb.exists():
        try:
            thumb_img = Image.open(thumb).convert("RGB")
            thumb_img.thumbnail((190, 190))
            x = thumb_box[0] + (190 - thumb_img.width) // 2
            y = thumb_box[1] + (190 - thumb_img.height) // 2
            image.paste(thumb_img, (x, y))
            draw.rectangle(thumb_box, outline=(230, 230, 230), width=2)
        except Exception:
            draw.rectangle(thumb_box, outline=(230, 230, 230), width=2)
    else:
        draw.rectangle(thumb_box, outline=(230, 230, 230), width=2)
        draw.text((55, 190), "sem imagem", fill=(120, 120, 120), font=small_font)

    x0 = 250
    y = 105
    name_lines = wrap(draw, clean(product.get("name")), name_font, width - x0 - 40)
    for line in name_lines[:2]:
        draw.text((x0, y), line, fill=(30, 30, 30), font=name_font)
        y += 38

    draw.text((x0, y + 8), card_text(product["permalink"]), fill=(0, 90, 170), font=mono_font)
    y += 48
    meta = f"Tags: {', '.join(product_tags(product)) or 'sem tag'} | Marcas: {', '.join(product_brands(product)) or 'sem marca'}"
    for line in wrap(draw, meta, small_font, width - x0 - 40)[:2]:
        draw.text((x0, y), line, fill=(85, 85, 85), font=small_font)
        y += 30

    cats = f"Categorias: {', '.join(product_categories(product)) or 'sem categoria'}"
    for line in wrap(draw, cats, small_font, width - x0 - 40)[:2]:
        draw.text((x0, y), line, fill=(85, 85, 85), font=small_font)
        y += 30

    y = 342
    draw.rectangle((30, y - 12, width - 30, y + 2), fill=(245, 245, 245))
    for issue in issue_items[:14]:
        text = f"{issue['kind']}: {issue['detail']}"
        if issue.get("expected"):
            text += f" | Corrigir: {issue['expected']}"
        lines = wrap(draw, text, body_font, width - 95)
        draw.text((45, y + 22), "- " + lines[0], fill=(160, 0, 0), font=body_font)
        y += 42
        for extra_line in lines[1:3]:
            draw.text((68, y + 12), extra_line, fill=(160, 0, 0), font=small_font)
            y += 31
        y += 8
    if len(issue_items) > 14:
        draw.text((45, y + 20), f"... mais {len(issue_items) - 14} erro(s) na tabela do documento", fill=(160, 0, 0), font=body_font)

    image.save(out, quality=86, optimize=True)
    return out


def make_non_product_card(title: str, link: str, detail: str, expected: str, index: int) -> Path:
    CARD_DIR.mkdir(parents=True, exist_ok=True)
    out = CARD_DIR / f"geral_{index:03d}.jpg"
    width, height = 1200, 430
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    header_font = font(36, True)
    name_font = font(31, True)
    body_font = font(28)
    mono_font = font(22)

    draw.rectangle((0, 0, width, 82), fill=(255, 105, 0))
    draw.text((28, 22), "Erro geral / menu / seguranca", fill="white", font=header_font)
    y = 115
    draw.text((36, y), card_text(title), fill=(30, 30, 30), font=name_font)
    y += 45
    if link:
        draw.text((36, y), card_text(link), fill=(0, 90, 170), font=mono_font)
        y += 42
    for label, text in [("Erro", detail), ("Corrigir", expected)]:
        lines = wrap(draw, f"{label}: {text}", body_font, width - 80)
        for line in lines[:3]:
            draw.text((36, y), line, fill=(160, 0, 0) if label == "Erro" else (40, 90, 40), font=body_font)
            y += 34
    image.save(out, quality=86, optimize=True)
    return out


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.color.rgb = RGBColor(35, 35, 35)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def issues_as_text(items: Iterable[dict]) -> str:
    return "\n".join(f"- {item['kind']}: {item['detail']}" for item in items)


def write_csv(products_by_id: dict[int, dict], issues: dict[int, list[dict]]):
    with OUT_CSV.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(["id", "produto", "link", "sku", "tags", "marcas", "categorias", "erros"])
        for product_id in sorted(issues):
            product = products_by_id[product_id]
            writer.writerow(
                [
                    product_id,
                    clean(product.get("name")),
                    product["permalink"],
                    clean(product.get("sku")),
                    ", ".join(product_tags(product)),
                    ", ".join(product_brands(product)),
                    ", ".join(product_categories(product)),
                    " | ".join(f"{item['kind']}: {item['detail']}" for item in issues[product_id]),
                ]
            )


def build_doc(products: list[dict], issues: dict[int, list[dict]], product_cards: dict[int, Path], general_cards: list[tuple[Path, str, str]]):
    products_by_id = {product["id"]: product for product in products}
    type_counter = Counter(item["kind"] for items in issues.values() for item in items)

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Relatorio COMPLETO Clickmed.pt - todos os erros detectados", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Data: 02/07/2026. Analise publica/passiva. Cada produto com erro tem link direto e imagem-evidencia.").italic = True

    doc.add_heading("Resumo", level=2)
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Item"
    summary.rows[0].cells[1].text = "Quantidade"
    for label, value in [
        ("Produtos analisados", str(len(products))),
        ("Produtos com pelo menos um erro/inconveniente", str(len(issues))),
        ("Total de erros/inconvenientes em produtos", str(sum(len(v) for v in issues.values()))),
        ("Erros gerais/menu/seguranca", str(len(general_cards))),
        ("Imagens-evidencia de produtos", str(len(product_cards))),
    ]:
        row = summary.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph("Contagem por tipo de erro:")
    counts = doc.add_table(rows=1, cols=2)
    counts.style = "Table Grid"
    counts.rows[0].cells[0].text = "Tipo"
    counts.rows[0].cells[1].text = "Quantidade"
    for kind, count in type_counter.most_common():
        row = counts.add_row().cells
        row[0].text = kind
        row[1].text = str(count)

    doc.add_heading("Tabela completa com links", level=2)
    doc.add_paragraph("Cada linha abaixo e um produto com erro. O link azul abre diretamente o produto.")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "ID"
    header[1].text = "Produto"
    header[2].text = "Link"
    header[3].text = "Erros"
    for product_id in sorted(issues):
        product = products_by_id[product_id]
        cells = table.add_row().cells
        cells[0].text = str(product_id)
        cells[1].text = clean(product.get("name"))
        add_hyperlink(cells[2].paragraphs[0], "Abrir produto", product["permalink"])
        cells[3].text = issues_as_text(issues[product_id])

    doc.add_page_break()
    doc.add_heading("Imagens-evidencia de TODOS os produtos com erro", level=2)
    doc.add_paragraph("Cada imagem abaixo resume os erros detectados naquele produto e mostra o link completo.")
    for index, product_id in enumerate(sorted(product_cards), 1):
        product = products_by_id[product_id]
        p = doc.add_paragraph()
        p.add_run(f"{index}. {clean(product.get('name'))}").bold = True
        p.add_run(" - ")
        add_hyperlink(p, "abrir produto", product["permalink"])
        doc.add_picture(str(product_cards[product_id]), width=Inches(6.85))

    doc.add_page_break()
    doc.add_heading("Erros gerais / categorias vazias / seguranca", level=2)
    for card_path, title_text, link in general_cards:
        p = doc.add_paragraph()
        p.add_run(title_text).bold = True
        if link:
            p.add_run(" - ")
            add_hyperlink(p, "abrir link", link)
        doc.add_picture(str(card_path), width=Inches(6.85))

    doc.save(OUT_DOCX)


def main():
    products = load_products()
    products_by_id = {product["id"]: product for product in products}
    issues = collect_product_issues(products)
    write_csv(products_by_id, issues)

    product_cards = {}
    total = len(issues)
    for i, product_id in enumerate(sorted(issues), 1):
        product_cards[product_id] = make_product_card(products_by_id[product_id], issues[product_id])
        if i % 50 == 0:
            print(f"cards {i}/{total}")

    general_cards: list[tuple[Path, str, str]] = []
    idx = 1
    for name, url in EMPTY_MENU_CATEGORIES:
        card = make_non_product_card(name, url, "Categoria aparece no menu, mas esta vazia/sem produtos", "Remover do menu ou esconder ate ter produtos", idx)
        general_cards.append((card, f"Categoria vazia: {name}", url))
        idx += 1
    for kind, detail, expected in SECURITY_FINDINGS:
        card = make_non_product_card(f"Seguranca - {kind}", "https://clickmed.pt/", detail, expected, idx)
        general_cards.append((card, f"Seguranca: {detail}", "https://clickmed.pt/"))
        idx += 1

    build_doc(products, issues, product_cards, general_cards)
    print(OUT_DOCX)
    print(OUT_CSV)
    print(CARD_DIR)


if __name__ == "__main__":
    main()
