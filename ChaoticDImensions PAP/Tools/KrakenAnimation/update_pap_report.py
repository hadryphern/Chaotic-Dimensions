from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


SECTION_TITLE = "Reconstrução da animação do Alien Kraken"


def format_body(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.6)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(12)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def detached_paragraph(document: Document, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    if style == "Normal":
        format_body(paragraph)
    element = paragraph._element
    document.element.body.remove(element)
    return element


def detached_picture(document: Document, image_path: Path):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run().add_picture(str(image_path), width=Cm(14.5))
    element = paragraph._element
    document.element.body.remove(element)
    return element


def insert_before(reference, elements) -> None:
    for element in elements:
        reference.addprevious(element)


def add_table_row(table, values: tuple[str, ...]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)


def find_table(document: Document, first_header: str):
    for table in document.tables:
        if table.rows and table.rows[0].cells:
            if table.rows[0].cells[0].text.strip() == first_header:
                return table
    raise RuntimeError(f"Tabela não encontrada: {first_header}")


def update_report(report_path: Path, preview_path: Path) -> None:
    document = Document(report_path)
    if any(paragraph.text.strip() == SECTION_TITLE for paragraph in document.paragraphs):
        print("O relatório já contém a secção da reconstrução do Kraken.")
        return

    legacy_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Conteúdo legado preservado"
    )

    section_elements = [
        detached_paragraph(document, SECTION_TITLE, "Heading 2"),
        detached_paragraph(
            document,
            (
                "A animação anterior do Alien Kraken apresentava diferenças visuais "
                "entre a cabeça e os tentáculos. A cabeça tinha contornos mais escuros, "
                "enquanto os membros inferiores pareciam elementos independentes. "
                "Alguns tentáculos centrais também mudavam de posição de forma brusca, "
                "o que criava a impressão de recortes deslocados."
            ),
            "Normal",
        ),
        detached_paragraph(
            document,
            (
                "Para corrigir o problema, as catorze camadas originais foram "
                "organizadas num rig bidimensional com raízes fixas. A deformação é "
                "aplicada progressivamente desde a base até à ponta de cada tentáculo, "
                "preservando a ligação ao corpo. Depois da composição é aplicado um "
                "único contorno azul-marinho, evitando diferenças entre a cabeça, o "
                "tronco e os membros."
            ),
            "Normal",
        ),
        detached_paragraph(
            document,
            (
                "Foram produzidos quatro atlas de 36 fotogramas: repouso, retorno, "
                "deslocamento horizontal e subida. O repouso e o retorno formam um "
                "ciclo contínuo de 72 fotogramas. A Ruby deixou de depender de uma "
                "lista de deslocamentos diferente para cada fotograma e passou a usar "
                "uma âncora única no corpo, também utilizada como origem dos lasers."
            ),
            "Normal",
        ),
        detached_picture(document, preview_path),
        detached_paragraph(
            document,
            (
                "Figura 7 - Comparação dos movimentos de repouso, retorno, "
                "deslocamento horizontal e subida do Alien Kraken."
            ),
            "PAP Note",
        ),
        detached_paragraph(
            document,
            (
                "A pipeline inclui testes de dimensões, transparência, continuidade "
                "do ciclo, estabilidade da cabeça, ligação das raízes e ausência de "
                "franjas verdes nas bordas. A compilação final terminou sem erros nem "
                "avisos do compilador."
            ),
            "Normal",
        ),
    ]
    insert_before(legacy_heading._element, section_elements)

    test_table = find_table(document, "Área testada")
    add_table_row(
        test_table,
        (
            "Animação do Alien Kraken",
            (
                "Executar 15 testes automáticos, validar os quatro atlas e compilar "
                "o projeto completo."
            ),
            (
                "72 fotogramas em ciclo, raízes ligadas, bordas limpas, Ruby estável "
                "e compilação sem erros."
            ),
        ),
    )

    error_table = find_table(document, "Problema")
    add_table_row(
        error_table,
        (
            "Tentáculos pareciam separados do corpo",
            (
                "As camadas eram deformadas sem uma raiz comum e o contorno variava "
                "entre partes da criatura."
            ),
            (
                "Criação de rig com raízes fixas, deformação progressiva e contorno "
                "unificado após a composição."
            ),
        ),
    )
    add_table_row(
        error_table,
        (
            "Erro CS1501 ao compilar a origem dos lasers",
            (
                "O método da âncora da Ruby deixou de receber o número do fotograma, "
                "mas o laser giratório ainda usava a assinatura anterior."
            ),
            (
                "A chamada foi atualizada para a âncora estável e o projeto voltou a "
                "compilar sem erros."
            ),
        ),
    )

    backup_dir = report_path.parent / "tmp" / "report_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup_path = backup_dir / f"{report_path.stem}_antes_animacao_kraken.docx"
    shutil.copy2(report_path, backup_path)

    temporary_path = report_path.with_name(f"{report_path.stem}.tmp.docx")
    document.save(temporary_path)
    temporary_path.replace(report_path)
    print(f"Relatório atualizado: {report_path}")
    print(f"Cópia de segurança: {backup_path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("report", type=Path)
    parser.add_argument("preview", type=Path)
    args = parser.parse_args()
    update_report(args.report.resolve(), args.preview.resolve())


if __name__ == "__main__":
    main()
