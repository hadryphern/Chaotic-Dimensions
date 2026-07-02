from __future__ import annotations

import shutil
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gerar_relatorio_blocos_final_docx import (
    REVIEW_REMOVED,
    confirmed_empty_categories,
    confirmed_security_findings,
    load_review_products,
    reviewed_issues,
)
from gerar_relatorio_completo_docx import add_hyperlink, clean, product_brands, product_categories, product_tags


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_SEM_IMAGENS_por_categoria.docx"
OUT_MD = ROOT / "relatorio_clickmed_SEM_IMAGENS_por_categoria.md"
OUT_VALIDATION = ROOT / "validacao_relatorio_clickmed_sem_imagens.md"

# Tambem atualiza o nome que o usuario ja estava usando.
LATEST_DOCX = ROOT / "relatorio_clickmed_FINAL_REVISADO.docx"
LATEST_MD = ROOT / "relatorio_clickmed_FINAL_REVISADO.md"
FINAL_BLOCK_DOCX = ROOT / "relatorio_clickmed_final_blocos.docx"
FINAL_BLOCK_MD = ROOT / "relatorio_clickmed_final_blocos.md"


TYPE_INFO = {
    "Marca": ("Alta", "Marca, icone ou filtro atribuido ao produto nao bate com o titulo.", "Corrigir marca/filtro/icone no produto."),
    "Tag/condicao": ("Alta", "Titulo diz Novo, Usado ou Open Box, mas a tag mostra outra condicao.", "Corrigir a tag de condicao."),
    "Tag/Grade": ("Alta", "Titulo informa Grade A/B/C, mas a tag nao informa a grade correta.", "Usar tag especifica, por exemplo Usado - Grade A."),
    "Tags": ("Media", "Produto nao tem tag de condicao.", "Adicionar Novo, Open Box ou Usado Grade A/B/C."),
    "Categoria": ("Alta", "Produto parece estar numa categoria errada de novo/usado.", "Mover para a categoria correta."),
    "Categorias": ("Media", "Produto esta sem categoria publica.", "Adicionar categoria correta."),
    "Descricao/IA": ("Alta", "Descricao tem resto claro de IA ou link colado indevido.", "Limpar e reescrever a descricao."),
    "Descricao/avaliacao": ("Alta", "Descricao fala em avaliacao/4.8, mas a API mostra zero reviews/rating.", "Remover avaliacao generica ou usar reviews reais."),
    "Descricao/campos vazios": ("Media", "Descricao tem campos vazios, como Bateria ou Versao.", "Preencher ou remover campos vazios."),
    "Descricao/imagem externa": ("Media", "Descricao usa imagem externa de outro site.", "Hospedar imagem propria ou remover."),
    "Descricao/especificacao": ("Alta", "Descricao menciona especificacao tecnica que parece nao bater com o produto.", "Corrigir especificacao tecnica."),
    "Descricao/Grade": ("Alta", "Grade no titulo e na descricao nao batem.", "Unificar grade no titulo e descricao."),
    "Descricao": ("Media", "Descricao esta ausente ou tem texto/template problemático.", "Rever descricao do produto."),
    "Imagem": ("Alta", "Alt/nome da imagem indica memoria/capacidade diferente do titulo.", "Corrigir imagem, alt ou nome do ficheiro."),
    "URL/slug": ("Media", "URL/slug indica capacidade diferente ou tem typo.", "Corrigir slug/URL."),
    "SKU": ("Baixa", "Produto esta sem SKU publico/API.", "Adicionar SKU/EAN/referencia interna."),
    "Duplicado": ("Media", "Existe outro produto com o mesmo nome.", "Unificar ou diferenciar estoque/unidade/localizacao."),
    "Promocao antiga": ("Media", "Produto continua numa categoria de campanha antiga.", "Remover de promocoes antigas ou atualizar campanha."),
    "Padronizacao": ("Baixa", "Titulo mistura idiomas/cores em ingles.", "Padronizar idioma do titulo se a loja quiser titulos em PT."),
}


GROUPS = [
    (
        "Marca, icone e filtros",
        "Erros que afetam filtros, icones de marca e a confianca do cliente ao procurar por marca.",
        ["Marca"],
    ),
    (
        "Tags, estado e grade",
        "Erros de Novo, Usado, Open Box e Grade A/B/C. Esta parte e importante porque muda a expectativa do cliente.",
        ["Tag/condicao", "Tag/Grade", "Tags"],
    ),
    (
        "Categorias dos produtos",
        "Produtos sem categoria ou colocados em categoria errada, principalmente novo/usado.",
        ["Categoria", "Categorias"],
    ),
    (
        "Descricao e conteudo",
        "Problemas no texto dos produtos: descricao vazia, campos incompletos, avaliacao falsa/generica, resto de IA ou especificacao errada.",
        [
            "Descricao/IA",
            "Descricao/avaliacao",
            "Descricao/campos vazios",
            "Descricao/imagem externa",
            "Descricao/especificacao",
            "Descricao/Grade",
            "Descricao",
        ],
    ),
    (
        "Imagem e URL",
        "Problemas em imagem, alt text, nome de ficheiro ou URL que podem confundir o cliente e o Google.",
        ["Imagem", "URL/slug"],
    ),
    (
        "Organizacao interna",
        "Problemas de gestao e manutencao: SKU, duplicados, promocoes antigas e padronizacao de titulos.",
        ["SKU", "Duplicado", "Promocao antiga", "Padronizacao"],
    ),
]


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(35, 35, 35)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def set_size(paragraph, size: float):
    for run in paragraph.runs:
        run.font.size = Pt(size)


def add_small_line(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run(label).bold = True
    p.add_run(text)
    set_size(p, 9.2)
    return p


def add_type_intro(doc: Document, kind: str, count: int):
    priority, meaning, fix = TYPE_INFO[kind]
    heading = doc.add_heading(f"{kind} ({count})", level=3)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Prioridade: ").bold = True
    p.add_run(priority)
    p.add_run(" | Significa: ").bold = True
    p.add_run(meaning)
    p.add_run(" | Corrigir: ").bold = True
    p.add_run(fix)
    set_size(p, 9.2)


def add_issue_block(doc: Document, product: dict, item: dict):
    title = clean(product.get("name"))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.add_run(f"{title} | ID {product['id']}").bold = True
    set_size(p, 10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run("Link: ").bold = True
    add_hyperlink(p, product["permalink"], product["permalink"])
    set_size(p, 8.8)

    add_small_line(doc, "Problema: ", item["detail"])
    if item.get("expected"):
        add_small_line(doc, "Corrigir: ", item["expected"])
    add_small_line(doc, "Tags / marcas / categorias: ", f"{', '.join(product_tags(product)) or 'sem tag'} | {', '.join(product_brands(product)) or 'sem marca'} | {', '.join(product_categories(product)) or 'sem categoria'}")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)


def build_indexes(products: list[dict], issues: dict[int, list[dict]]):
    products_by_id = {product["id"]: product for product in products}
    by_kind: dict[str, list[tuple[dict, dict]]] = defaultdict(list)
    for product_id, items in issues.items():
        product = products_by_id[product_id]
        for item in items:
            by_kind[item["kind"]].append((product, item))
    for kind in by_kind:
        by_kind[kind].sort(key=lambda row: (clean(row[0].get("name")).lower(), row[0]["id"]))
    return by_kind


def add_summary(doc: Document, products: list[dict], issues: dict[int, list[dict]], by_kind: dict[str, list[tuple[dict, dict]]], empty_categories, security):
    doc.add_heading("Resumo", level=2)
    summary_lines = [
        f"Produtos analisados: {len(products)}",
        f"Produtos com pelo menos um problema: {len(issues)}",
        f"Problemas de produto depois da revisao: {sum(len(items) for items in issues.values())}",
        f"Categorias vazias confirmadas: {len(empty_categories)}",
        f"Pontos de seguranca/configuracao confirmados: {len(security)}",
        f"Itens removidos para evitar falso positivo: {len(REVIEW_REMOVED)}",
        "Imagens removidas deste DOCX: sim, o documento nao contem imagens.",
    ]
    for line in summary_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Contagem por categoria", level=2)
    for group_title, _, kinds in GROUPS:
        count = sum(len(by_kind.get(kind, [])) for kind in kinds)
        doc.add_paragraph(f"{group_title}: {count}", style="List Bullet")
    doc.add_paragraph(f"Menu/categorias vazias: {len(empty_categories)}", style="List Bullet")
    doc.add_paragraph(f"Seguranca/configuracao: {len(security)}", style="List Bullet")


def add_general_sections(doc: Document, empty_categories, security):
    doc.add_page_break()
    doc.add_heading("Menu e categorias vazias", level=2)
    doc.add_paragraph("Categorias que aparecem no menu, mas abrem pagina sem produtos.")
    for name, url in empty_categories:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run(name).bold = True
        p.add_run(" | Link: ")
        add_hyperlink(p, url, url)
        set_size(p, 9.2)
        add_small_line(doc, "Problema: ", "Categoria aparece no menu e esta vazia/sem produtos.")
        add_small_line(doc, "Corrigir: ", "Remover do menu ou esconder ate ter produtos.")

    doc.add_page_break()
    doc.add_heading("Seguranca e configuracao", level=2)
    doc.add_paragraph("Pontos confirmados por consulta publica/passiva. Nao e teste invasivo.")
    for title, detail, expected, url in security:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run(title).bold = True
        p.add_run(" | Link: ")
        add_hyperlink(p, url, url)
        set_size(p, 9.2)
        add_small_line(doc, "Problema: ", detail)
        add_small_line(doc, "Corrigir: ", expected)


def build_doc(products: list[dict], issues: dict[int, list[dict]], empty_categories, security):
    by_kind = build_indexes(products, issues)
    kind_counter = Counter(kind for kind, rows in by_kind.items() for _ in rows)

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Clickmed.pt - relatorio sem imagens por categoria", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Versao sem imagens, sem tabelas, separada por tipo de erro.").italic = True

    add_summary(doc, products, issues, by_kind, empty_categories, security)

    doc.add_heading("Tipos de erro encontrados", level=2)
    for kind, count in kind_counter.most_common():
        doc.add_paragraph(f"{kind}: {count}", style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Problemas de produtos por categoria", level=2)
    doc.add_paragraph("O mesmo produto pode aparecer em mais de uma categoria quando tem mais de um tipo de problema.")

    for group_title, group_desc, kinds in GROUPS:
        group_total = sum(len(by_kind.get(kind, [])) for kind in kinds)
        if group_total == 0:
            continue
        doc.add_page_break()
        doc.add_heading(f"{group_title} ({group_total})", level=2)
        doc.add_paragraph(group_desc)
        for kind in kinds:
            rows = by_kind.get(kind, [])
            if not rows:
                continue
            add_type_intro(doc, kind, len(rows))
            for product, item in rows:
                add_issue_block(doc, product, item)

    add_general_sections(doc, empty_categories, security)
    doc.save(OUT_DOCX)
    shutil.copyfile(OUT_DOCX, LATEST_DOCX)
    shutil.copyfile(OUT_DOCX, FINAL_BLOCK_DOCX)


def build_md(products: list[dict], issues: dict[int, list[dict]], empty_categories, security):
    by_kind = build_indexes(products, issues)
    lines = [
        "# Clickmed.pt - relatorio sem imagens por categoria",
        "",
        "Versao sem imagens, sem tabelas, separada por tipo de erro.",
        "",
        "## Resumo",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com pelo menos um problema: {len(issues)}",
        f"- Problemas de produto depois da revisao: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias confirmadas: {len(empty_categories)}",
        f"- Pontos de seguranca/configuracao confirmados: {len(security)}",
        f"- Itens removidos para evitar falso positivo: {len(REVIEW_REMOVED)}",
        f"- Imagens removidas deste DOCX: sim",
        "",
    ]
    for group_title, group_desc, kinds in GROUPS:
        group_total = sum(len(by_kind.get(kind, [])) for kind in kinds)
        if group_total == 0:
            continue
        lines += [f"## {group_title} ({group_total})", "", group_desc, ""]
        for kind in kinds:
            rows = by_kind.get(kind, [])
            if not rows:
                continue
            priority, meaning, fix = TYPE_INFO[kind]
            lines += [f"### {kind} ({len(rows)})", "", f"Prioridade: {priority}", f"Significa: {meaning}", f"Corrigir: {fix}", ""]
            for product, item in rows:
                lines.append(f"#### {clean(product.get('name'))} | ID {product['id']}")
                lines.append(f"Link: {product['permalink']}")
                lines.append(f"Problema: {item['detail']}")
                if item.get("expected"):
                    lines.append(f"Corrigir: {item['expected']}")
                lines.append(f"Tags / marcas / categorias: {', '.join(product_tags(product)) or 'sem tag'} | {', '.join(product_brands(product)) or 'sem marca'} | {', '.join(product_categories(product)) or 'sem categoria'}")
                lines.append("")
        lines.append("")

    lines += ["## Menu e categorias vazias", ""]
    for name, url in empty_categories:
        lines += [f"### {name}", f"Link: {url}", "Problema: Categoria aparece no menu e esta vazia/sem produtos.", "Corrigir: Remover do menu ou esconder ate ter produtos.", ""]

    lines += ["## Seguranca e configuracao", "", "Pontos confirmados por consulta publica/passiva. Nao e teste invasivo.", ""]
    for title, detail, expected, url in security:
        lines += [f"### {title}", f"Link: {url}", f"Problema: {detail}", f"Corrigir: {expected}", ""]

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    shutil.copyfile(OUT_MD, LATEST_MD)
    shutil.copyfile(OUT_MD, FINAL_BLOCK_MD)


def validate(products: list[dict], issues: dict[int, list[dict]], empty_categories, security):
    doc = Document(OUT_DOCX)
    with zipfile.ZipFile(OUT_DOCX) as archive:
        bad = archive.testzip()
    lines = [
        "# Validacao do relatorio sem imagens",
        "",
        f"- DOCX integro: {'sim' if bad is None else 'nao'}",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com problemas: {len(issues)}",
        f"- Problemas de produto: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias confirmadas: {len(empty_categories)}",
        f"- Seguranca/configuracao confirmados: {len(security)}",
        f"- Tabelas no DOCX: {len(doc.tables)}",
        f"- Imagens no DOCX: {len(doc.inline_shapes)}",
        f"- Paragrafos no DOCX: {len(doc.paragraphs)}",
        f"- Arquivo atualizado tambem em relatorio_clickmed_FINAL_REVISADO.docx: sim",
        f"- Arquivo atualizado tambem em relatorio_clickmed_final_blocos.docx: sim",
    ]
    OUT_VALIDATION.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_review_products()
    issues = reviewed_issues(products)
    empty_categories = confirmed_empty_categories()
    security = confirmed_security_findings()
    build_doc(products, issues, empty_categories, security)
    build_md(products, issues, empty_categories, security)
    validate(products, issues, empty_categories, security)
    print(OUT_DOCX)
    print(OUT_MD)
    print(OUT_VALIDATION)


if __name__ == "__main__":
    main()
