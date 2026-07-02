from __future__ import annotations

import json
import re
import shutil
import urllib.error
import urllib.request
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import gerar_relatorio_completo_docx as core
from gerar_relatorio_completo_docx import (
    EMPTY_MENU_CATEGORIES,
    add_hyperlink,
    clean,
    collect_product_issues,
    detect_brand,
    load_products,
    make_non_product_card,
    make_product_card,
    product_brands,
    product_categories,
    product_tags,
)


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_final_blocos.docx"
OUT_MD = ROOT / "relatorio_clickmed_final_blocos.md"
OUT_DOCX_REVIEWED = ROOT / "relatorio_clickmed_FINAL_REVISADO.docx"
OUT_MD_REVIEWED = ROOT / "relatorio_clickmed_FINAL_REVISADO.md"
VALIDATION_MD = ROOT / "validacao_relatorio_clickmed_final.md"
FRESH_PRODUCTS = ROOT / "clickmed_products_fresh.json"
REVIEW_CARD_DIR = ROOT / "evidencias_clickmed_revisado"
EXISTING_THUMB_DIR = ROOT / "evidencias_clickmed_todos_produtos" / "_thumbs"
IMAGE_WIDTH = Inches(5.85)

CONFIRMED_OLD_PROMO_CATS = {"Promocoes Marco", "Promoções Março", "Promoções Abril", "Promoções Maio"}
COMPATIBLE_ACCESSORY_TERMS = [
    "compatível",
    "compativel",
    "bracelet",
    "pulseira",
    "capa",
    "case",
    "pelicula",
    "película",
    "vidro",
    "bateria",
    "comando",
    "adaptador",
]

REVIEW_REMOVED: list[tuple[int, str, str]] = []


PRIORITY = [
    "Marca",
    "Tag/condicao",
    "Tag/Grade",
    "Categoria",
    "Imagem",
    "URL/slug",
    "Descricao/avaliacao",
    "Descricao/especificacao",
    "Descricao/IA",
    "Descricao",
    "Descricao/campos vazios",
    "Descricao/imagem externa",
    "Tags",
    "Categorias",
    "Duplicado",
    "SKU",
    "Promocao antiga",
    "Padronizacao",
]


SECTION_TITLES = {
    "critico": "Produtos com erro mais grave",
    "conteudo": "Produtos com erro de descricao, imagem ou URL",
    "organizacao": "Produtos com erro de organizacao",
}


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.color.rgb = RGBColor(35, 35, 35)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)


def keep_same_page(paragraphs):
    for paragraph in paragraphs:
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.widow_control = True
    for paragraph in paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True
    if paragraphs:
        paragraphs[-1].paragraph_format.keep_with_next = False


def set_runs_size(paragraph, size_pt: float):
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)


def section_for_items(items: list[dict]) -> str:
    kinds = {item["kind"] for item in items}
    if kinds & {"Marca", "Tag/condicao", "Tag/Grade", "Categoria", "Imagem", "URL/slug"}:
        return "critico"
    if any(kind.startswith("Descricao") for kind in kinds):
        return "conteudo"
    return "organizacao"


def sort_key(product: dict, items: list[dict]):
    kinds = [item["kind"] for item in items]
    best = min((PRIORITY.index(kind) if kind in PRIORITY else 999 for kind in kinds), default=999)
    return best, clean(product.get("name")).lower()


def load_review_products() -> list[dict]:
    if FRESH_PRODUCTS.exists():
        return json.loads(FRESH_PRODUCTS.read_text(encoding="utf-8"))
    return load_products()


def should_drop_compatible_brand_issue(product: dict, item: dict) -> bool:
    if item["kind"] != "Marca":
        return False
    brands = product_brands(product)
    if brands != ["Compatível"]:
        return False
    name = clean(product.get("name")).lower()
    return any(term in name for term in COMPATIBLE_ACCESSORY_TERMS)


def review_issue(product: dict, item: dict) -> dict | None:
    item = dict(item)

    if should_drop_compatible_brand_issue(product, item):
        REVIEW_REMOVED.append((product["id"], item["kind"], "marca compativel/compatibilidade, baixo risco de falso positivo"))
        return None

    if item["kind"] == "Marca":
        brands = product_brands(product)
        title_brands = sorted(detect_brand(clean(product.get("name"))))
        if brands == ["Compatível"] and title_brands:
            item["detail"] = (
                f"Marca/filtro a confirmar: o titulo indica {', '.join(title_brands)}, "
                "mas a marca atribuida e apenas Compativel"
            )
            item["expected"] = "Confirmar se e produto original ou acessorio compativel; ajustar marca/filtro se necessario"
        elif title_brands and brands:
            item["detail"] = f"Marca/filtro incorreto: titulo indica {', '.join(title_brands)}, mas marca atribuida e {', '.join(brands)}"
            item["expected"] = "Corrigir marca/filtro/icone"

    if item["kind"] == "Promocao antiga":
        old_promos = [cat for cat in product_categories(product) if cat in CONFIRMED_OLD_PROMO_CATS]
        if not old_promos:
            REVIEW_REMOVED.append((product["id"], item["kind"], "campanha Junho removida para evitar falso positivo"))
            return None
        item["detail"] = f"Produto ainda em categoria de campanha antiga: {', '.join(old_promos)}"
        item["expected"] = "Remover de promocoes antigas ou atualizar campanha"

    if item["kind"] == "Padronizacao":
        item["detail"] = "Melhoria de padronizacao: " + item["detail"]
        item["expected"] = "Padronizar idioma do titulo se a loja quiser manter todos os titulos em PT"

    if item["kind"] == "SKU":
        item["detail"] = "Gestao interna: produto sem SKU publico/API"
        item["expected"] = "Adicionar SKU/EAN/referencia interna"

    return item


def reviewed_issues(products: list[dict]) -> dict[int, list[dict]]:
    REVIEW_REMOVED.clear()
    raw = collect_product_issues(products)
    by_id = {product["id"]: product for product in products}
    revised: dict[int, list[dict]] = {}
    for product_id, items in raw.items():
        product = by_id[product_id]
        kept = []
        for item in items:
            reviewed = review_issue(product, item)
            if reviewed is not None and reviewed not in kept:
                kept.append(reviewed)
        if kept:
            revised[product_id] = kept
    return revised


def add_divider(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    p.add_run(" ")


def add_product_block(doc: Document, product: dict, items: list[dict], image_path: Path):
    block_paragraphs = []

    heading = doc.add_heading(f"{clean(product.get('name'))} | ID {product['id']}", level=3)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(4)
    block_paragraphs.append(heading)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Link: ").bold = True
    add_hyperlink(p, product["permalink"], product["permalink"])
    set_runs_size(p, 9)
    block_paragraphs.append(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Problemas encontrados: ").bold = True
    p.add_run(str(len(items)))
    for item in items:
        p.add_run().add_break()
        p.add_run("- ").bold = True
        p.add_run(f"{item['kind']}: ").bold = True
        p.add_run(item["detail"])
        if item.get("expected"):
            p.add_run(" | Corrigir: ").bold = True
            p.add_run(item["expected"])
    set_runs_size(p, 8.7)
    block_paragraphs.append(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Tags: ").bold = True
    p.add_run(", ".join(product_tags(product)) or "sem tag")
    p.add_run(" | Marcas: ").bold = True
    p.add_run(", ".join(product_brands(product)) or "sem marca")
    set_runs_size(p, 9)
    block_paragraphs.append(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run("Categorias: ").bold = True
    p.add_run(", ".join(product_categories(product)) or "sem categoria")
    set_runs_size(p, 9)
    block_paragraphs.append(p)

    doc.add_picture(str(image_path), width=IMAGE_WIDTH)
    image_paragraph = doc.paragraphs[-1]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_after = Pt(18)
    block_paragraphs.append(image_paragraph)

    keep_same_page(block_paragraphs)
    add_divider(doc)


def make_cards(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, Path]:
    core.CARD_DIR = REVIEW_CARD_DIR
    core.THUMB_DIR = EXISTING_THUMB_DIR if EXISTING_THUMB_DIR.exists() else REVIEW_CARD_DIR / "_thumbs"
    REVIEW_CARD_DIR.mkdir(parents=True, exist_ok=True)
    products_by_id = {product["id"]: product for product in products}
    cards = {}
    for product_id in sorted(issues):
        old_card = REVIEW_CARD_DIR / f"produto_{product_id}.jpg"
        if old_card.exists():
            old_card.unlink()
        cards[product_id] = make_product_card(products_by_id[product_id], issues[product_id])
    return cards


def url_text(url: str) -> tuple[int | None, str, dict[str, str]]:
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=20) as response:
            return response.status, response.read().decode("utf-8", "replace"), dict(response.headers)
    except urllib.error.HTTPError as error:
        body = error.read().decode("utf-8", "replace")
        return error.code, body, dict(error.headers)
    except Exception:
        return None, "", {}


def confirmed_empty_categories() -> list[tuple[str, str]]:
    status, home, _ = url_text("https://clickmed.pt/")
    confirmed = []
    for name, url in EMPTY_MENU_CATEGORIES:
        category_status, body, _ = url_text(url)
        appears_in_menu = url.replace("https://clickmed.pt", "") in home or url in home
        no_products = "Não foram encontrados produtos" in body or "No products were found" in body
        if status == 200 and category_status == 200 and appears_in_menu and no_products:
            confirmed.append((name, url))
    return confirmed


def confirmed_security_findings() -> list[tuple[str, str, str, str]]:
    findings: list[tuple[str, str, str, str]] = []
    status, home, headers = url_text("https://clickmed.pt/")
    lower_headers = {key.lower(): value for key, value in headers.items()}
    if status == 200:
        if "strict-transport-security" not in lower_headers:
            findings.append(("Seguranca - HSTS", "Homepage sem Strict-Transport-Security visivel", "Ativar HSTS em todas as paginas HTTPS", "https://clickmed.pt/"))
        if "content-security-policy" not in lower_headers:
            findings.append(("Seguranca - CSP", "Homepage sem Content-Security-Policy geral visivel", "Definir CSP e pelo menos frame-ancestors 'self'", "https://clickmed.pt/"))
        if "referrer-policy" not in lower_headers:
            findings.append(("Seguranca - Referrer-Policy", "Homepage sem Referrer-Policy visivel", "Usar Referrer-Policy: strict-origin-when-cross-origin", "https://clickmed.pt/"))
        if "permissions-policy" not in lower_headers:
            findings.append(("Seguranca - Permissions-Policy", "Homepage sem Permissions-Policy visivel", "Bloquear camera/microfone/geolocalizacao se nao forem usados", "https://clickmed.pt/"))
        if "author/admin" in home or 'content="admin"' in home:
            findings.append(("Seguranca - autor admin", "HTML/JSON-LD expoe autor chamado admin", "Evitar usuario publico chamado admin e rever autor das paginas", "https://clickmed.pt/"))
        if re.search(r"wp-content/(plugins|themes)/[^\"']+\?ver=", home):
            findings.append(("Seguranca - versoes expostas", "HTML expoe versoes de tema/plugins em varios assets", "Manter tudo atualizado e reduzir exposicao quando possivel", "https://clickmed.pt/"))

    xml_status, _, xml_headers = url_text("https://clickmed.pt/xmlrpc.php")
    if xml_status in {200, 405} or "POST" in xml_headers.get("Allow", ""):
        findings.append(("Seguranca - XML-RPC", "xmlrpc.php responde publicamente", "Desativar XML-RPC se nao for usado ou proteger com WAF/rate limit", "https://clickmed.pt/xmlrpc.php"))

    wp_status, wp_json, _ = url_text("https://clickmed.pt/wp-json/")
    if wp_status == 200:
        exposed = [name for name in ["sequra", "klarna", "nakedcat-recommend-ifthenpay", "mwai", "mcp", "wc/private"] if name in wp_json]
        if exposed:
            findings.append(("Seguranca - REST API", f"REST API publica mostra rotas/plugins sensiveis: {', '.join(exposed)}", "Desativar rotas/plugins nao usados e garantir auth/rate-limit", "https://clickmed.pt/wp-json/"))
        if any(name in wp_json for name in ["sequra", "klarna", "ifthenpay"]):
            findings.append(("Seguranca - pagamentos", "Rotas de gateways/pagamentos aparecem no indice REST publico", "Confirmar assinatura/secret nos webhooks e nunca confiar em valor vindo do front-end", "https://clickmed.pt/wp-json/"))

    login_status, _, _ = url_text("https://clickmed.pt/wp-login.php")
    if login_status == 200:
        findings.append(("Seguranca - wp-login", "wp-login.php esta publico", "Ativar 2FA, WAF/rate-limit, senha forte e menor privilegio possivel", "https://clickmed.pt/wp-login.php"))

    return findings


def make_general_cards():
    core.CARD_DIR = REVIEW_CARD_DIR
    core.THUMB_DIR = EXISTING_THUMB_DIR if EXISTING_THUMB_DIR.exists() else REVIEW_CARD_DIR / "_thumbs"
    cards = []
    index = 1
    for name, url in confirmed_empty_categories():
        title = f"Categoria vazia - {name}"
        detail = "Categoria aparece no menu e abre pagina sem produtos."
        expected = "Remover do menu ou esconder ate ter produtos."
        card = make_non_product_card(title, url, detail, expected, index)
        cards.append((card, title, f"{detail} Corrigir: {expected}", url))
        index += 1
    for title, detail, expected, url in confirmed_security_findings():
        card = make_non_product_card(title, url, detail, expected, index)
        cards.append((card, title, f"{detail} Corrigir: {expected}", url))
        index += 1
    return cards


def build_doc(products: list[dict], issues: dict[int, list[dict]]):
    products_by_id = {product["id"]: product for product in products}
    cards = make_cards(products, issues)
    general_cards = make_general_cards()

    grouped: dict[str, list[tuple[dict, list[dict]]]] = {key: [] for key in SECTION_TITLES}
    for product_id, items in issues.items():
        product = products_by_id[product_id]
        grouped[section_for_items(items)].append((product, items))
    for key in grouped:
        grouped[key].sort(key=lambda row: sort_key(row[0], row[1]))

    kind_counter = Counter(item["kind"] for items in issues.values() for item in items)

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Clickmed.pt - problemas encontrados", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Versao final em blocos: nome, problema, link e imagem juntos. Sem tabelas.").italic = True

    doc.add_heading("Resumo", level=2)
    for text in [
        f"Produtos analisados: {len(products)}",
        f"Produtos com erro/inconveniente: {len(issues)}",
        f"Total de problemas encontrados: {sum(len(items) for items in issues.values())}",
        f"Categorias vazias no menu confirmadas: {sum(1 for _, title, _, _ in general_cards if title.startswith('Categoria vazia'))}",
        f"Pontos de seguranca/configuracao confirmados: {sum(1 for _, title, _, _ in general_cards if title.startswith('Seguranca'))}",
        f"Itens removidos na revisao para evitar falso positivo: {len(REVIEW_REMOVED)}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Tipos de erro encontrados", level=2)
    for kind, count in kind_counter.most_common():
        doc.add_paragraph(f"{kind}: {count}", style="List Bullet")

    doc.add_paragraph(
        "Cada produto aparece apenas uma vez. Dentro do bloco ficam todos os problemas encontrados nesse produto, o link completo e a imagem-evidencia."
    )

    doc.add_page_break()

    for key, title_text in SECTION_TITLES.items():
        rows = grouped[key]
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(f"Produtos nesta parte: {len(rows)}")
        for product, items in rows:
            add_product_block(doc, product, items, cards[product["id"]])
        doc.add_page_break()

    doc.add_heading("Categorias vazias no menu e seguranca", level=2)
    for card, title_text, detail, link in general_cards:
        block_paragraphs = []
        heading = doc.add_heading(title_text, level=3)
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after = Pt(4)
        block_paragraphs.append(heading)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run("Problema: ").bold = True
        p.add_run(detail)
        set_runs_size(p, 9)
        block_paragraphs.append(p)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.add_run("Link: ").bold = True
        add_hyperlink(p, link, link)
        set_runs_size(p, 9)
        block_paragraphs.append(p)

        doc.add_picture(str(card), width=IMAGE_WIDTH)
        image_paragraph = doc.paragraphs[-1]
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(18)
        block_paragraphs.append(image_paragraph)

        keep_same_page(block_paragraphs)
        add_divider(doc)

    doc.save(OUT_DOCX_REVIEWED)
    shutil.copyfile(OUT_DOCX_REVIEWED, OUT_DOCX)
    return general_cards


def build_md(products: list[dict], issues: dict[int, list[dict]], general_cards: list[tuple[Path, str, str, str]]):
    products_by_id = {product["id"]: product for product in products}
    grouped = {key: [] for key in SECTION_TITLES}
    for product_id, items in issues.items():
        product = products_by_id[product_id]
        grouped[section_for_items(items)].append((product, items))
    for key in grouped:
        grouped[key].sort(key=lambda row: sort_key(row[0], row[1]))

    lines = [
        "# Clickmed.pt - problemas encontrados",
        "",
        "Versao final em blocos: nome, problema, link e imagem juntos. Sem tabelas.",
        "",
        "## Resumo",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com erro/inconveniente: {len(issues)}",
        f"- Total de problemas encontrados: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias no menu confirmadas: {sum(1 for _, title, _, _ in general_cards if title.startswith('Categoria vazia'))}",
        f"- Pontos de seguranca/configuracao confirmados: {sum(1 for _, title, _, _ in general_cards if title.startswith('Seguranca'))}",
        f"- Itens removidos na revisao para evitar falso positivo: {len(REVIEW_REMOVED)}",
        "",
    ]
    for key, title_text in SECTION_TITLES.items():
        lines += [f"## {title_text}", "", f"Produtos nesta parte: {len(grouped[key])}", ""]
        for product, items in grouped[key]:
            lines.append(f"### {clean(product.get('name'))} | ID {product['id']}")
            lines.append("")
            lines.append(f"Link: {product['permalink']}")
            lines.append("")
            lines.append(f"Tags: {', '.join(product_tags(product)) or 'sem tag'}")
            lines.append(f"Marcas: {', '.join(product_brands(product)) or 'sem marca'}")
            lines.append(f"Categorias: {', '.join(product_categories(product)) or 'sem categoria'}")
            lines.append("")
            lines.append("Problemas:")
            for item in items:
                line = f"- {item['kind']}: {item['detail']}"
                if item.get("expected"):
                    line += f" | Corrigir: {item['expected']}"
                lines.append(line)
            lines.append("")
            lines.append("---")
            lines.append("")
    OUT_MD_REVIEWED.write_text("\n".join(lines), encoding="utf-8")
    shutil.copyfile(OUT_MD_REVIEWED, OUT_MD)


def paragraph_has_image(paragraph) -> bool:
    return "<w:drawing>" in paragraph._p.xml


def validate_doc(products: list[dict], issues: dict[int, list[dict]], general_cards: list[tuple[Path, str, str, str]]):
    doc = Document(OUT_DOCX_REVIEWED)
    heading3 = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.style.name == "Heading 3"]
    product_titles = [f"{clean(product.get('name'))} | ID {product['id']}" for product in products if product["id"] in issues]
    duplicate_product_names = Counter(clean(product.get("name")).lower() for product in products)
    duplicate_name_count = sum(1 for name, count in duplicate_product_names.items() if count > 1)

    blocks_missing_image = []
    blocks_missing_link = []
    current_heading = None
    current_has_image = False
    current_has_link = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 3":
            if current_heading is not None:
                if not current_has_image:
                    blocks_missing_image.append(current_heading)
                if not current_has_link:
                    blocks_missing_link.append(current_heading)
            current_heading = paragraph.text.strip()
            current_has_image = False
            current_has_link = False
        elif current_heading is not None:
            if paragraph_has_image(paragraph):
                current_has_image = True
            if "https://clickmed.pt/" in paragraph.text:
                current_has_link = True
    if current_heading is not None:
        if not current_has_image:
            blocks_missing_image.append(current_heading)
        if not current_has_link:
            blocks_missing_link.append(current_heading)

    with zipfile.ZipFile(OUT_DOCX_REVIEWED) as archive:
        archive.testzip()

    lines = [
        "# Validacao final do relatorio Clickmed",
        "",
        f"- Produtos carregados da fonte atual: {len(products)}",
        f"- Produtos com problemas apos revisao conservadora: {len(issues)}",
        f"- Total de problemas apos revisao conservadora: {sum(len(items) for items in issues.values())}",
        f"- Blocos gerais confirmados: {len(general_cards)}",
        f"- Tabelas no DOCX: {len(doc.tables)}",
        f"- Imagens no DOCX: {len(doc.inline_shapes)}",
        f"- Imagens esperadas: {len(issues) + len(general_cards)}",
        f"- Titulos de produto esperados: {len(product_titles)}",
        f"- Titulos Heading 3 no documento: {len(heading3)}",
        f"- Titulos Heading 3 duplicados: {sum(1 for _, count in Counter(heading3).items() if count > 1)}",
        f"- Blocos sem imagem detectados: {len(blocks_missing_image)}",
        f"- Blocos sem link completo detectados: {len(blocks_missing_link)}",
        f"- Nomes realmente duplicados no site/API: {duplicate_name_count}",
        f"- Itens removidos para evitar falso positivo: {len(REVIEW_REMOVED)}",
        "",
        "Observacao: nomes duplicados no site foram mantidos como problema real, mas os titulos do relatorio agora incluem ID para nao parecer duplicacao acidental do DOCX.",
        "",
        "Itens removidos na revisao:",
    ]
    if REVIEW_REMOVED:
        for product_id, kind, reason in REVIEW_REMOVED[:200]:
            lines.append(f"- ID {product_id} | {kind} | {reason}")
    else:
        lines.append("- Nenhum.")

    VALIDATION_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_review_products()
    issues = reviewed_issues(products)
    general_cards = build_doc(products, issues)
    build_md(products, issues, general_cards)
    validate_doc(products, issues, general_cards)
    print(OUT_DOCX_REVIEWED)
    print(OUT_MD_REVIEWED)
    print(VALIDATION_MD)


if __name__ == "__main__":
    main()
