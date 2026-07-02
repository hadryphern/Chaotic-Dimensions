from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gerar_relatorio_completo_docx import (
    EMPTY_MENU_CATEGORIES,
    SECURITY_FINDINGS,
    add_hyperlink,
    clean,
    collect_product_issues,
    load_products,
    make_non_product_card,
    make_product_card,
    product_brands,
    product_categories,
    product_tags,
)
from gerar_relatorio_humanizado_docx import GROUPS, GROUP_TEXT, group_for_kind, grouped_rows


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_sem_tabelas.docx"
OUT_MD = ROOT / "relatorio_clickmed_sem_tabelas.md"


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.color.rgb = RGBColor(35, 35, 35)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def bullet(doc: Document, text: str = "", bold_start: str | None = None):
    paragraph = doc.add_paragraph(style="List Bullet")
    if bold_start:
        run = paragraph.add_run(bold_start)
        run.bold = True
        if text:
            paragraph.add_run(text)
    else:
        paragraph.add_run(text)
    return paragraph


def add_product_block(doc: Document, product: dict, items: list[dict]):
    title = clean(product.get("name"))
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.add_run(title).bold = True
    p.add_run(" | ")
    add_hyperlink(p, "abrir produto", product["permalink"])

    meta = []
    tags = ", ".join(product_tags(product)) or "sem tag"
    brands = ", ".join(product_brands(product)) or "sem marca"
    cats = ", ".join(product_categories(product)) or "sem categoria"
    meta.append(f"Tags: {tags}")
    meta.append(f"Marcas: {brands}")
    meta.append(f"Categorias: {cats}")
    for line in meta:
        bullet(doc, line)

    for item in items:
        problem = f"{item['kind']}: {item['detail']}"
        if item.get("expected"):
            problem += f" | Corrigir: {item['expected']}"
        bullet(doc, problem)


def make_cards(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, Path]:
    products_by_id = {product["id"]: product for product in products}
    cards = {}
    for product_id in sorted(issues):
        cards[product_id] = make_product_card(products_by_id[product_id], issues[product_id])
    return cards


def make_general_cards() -> list[tuple[Path, str, str]]:
    cards = []
    index = 1
    for name, url in EMPTY_MENU_CATEGORIES:
        card = make_non_product_card(
            name,
            url,
            "Categoria aparece no menu, mas esta vazia/sem produtos.",
            "Remover do menu ou esconder ate ter produtos.",
            index,
        )
        cards.append((card, f"Categoria vazia: {name}", url))
        index += 1

    for kind, detail, expected in SECURITY_FINDINGS:
        card = make_non_product_card(f"Seguranca - {kind}", "https://clickmed.pt/", detail, expected, index)
        cards.append((card, f"Seguranca: {detail}", "https://clickmed.pt/"))
        index += 1
    return cards


def build_doc(products: list[dict], issues: dict[int, list[dict]]):
    grouped = grouped_rows(products, issues)
    cards = make_cards(products, issues)
    general_cards = make_general_cards()

    total_issues = sum(len(items) for items in issues.values())
    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Clickmed.pt - relatorio simples de correcoes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Versao sem tabelas. Feito para leitura e correcao rapida.").italic = True

    doc.add_heading("Resumo direto", level=2)
    bullet(doc, f"Foram analisados {len(products)} produtos publicos.")
    bullet(doc, f"{len(issues)} produtos tem pelo menos um erro ou ponto incomodo para corrigir.")
    bullet(doc, f"No total foram encontrados {total_issues} problemas nos produtos.")
    bullet(doc, f"Alem disso, ha {len(EMPTY_MENU_CATEGORIES)} categorias vazias no menu e {len(SECURITY_FINDINGS)} pontos de seguranca/configuracao.")
    bullet(doc, "O link de cada produto esta logo ao lado do nome, para abrir sem procurar manualmente.")
    bullet(doc, "As imagens-evidencia ficam no fim do documento, uma por produto com erro.")

    doc.add_heading("Ordem sugerida para corrigir", level=2)
    bullet(doc, "1. Corrigir tags de estado: Novo, Open Box, Usado Grade A, B ou C.")
    bullet(doc, "2. Corrigir marcas/filtros, porque isso afeta icones e pesquisa.")
    bullet(doc, "3. Remover descricoes com avaliacao falsa, restos de IA, campos vazios e especificacoes erradas.")
    bullet(doc, "4. Rever imagens, alt text e URLs que mostram capacidade/modelo errado.")
    bullet(doc, "5. Limpar categorias vazias, promocoes antigas, SKUs ausentes e produtos duplicados.")

    doc.add_page_break()

    for group in GROUPS:
        rows = grouped.get(group, [])
        text = GROUP_TEXT[group]
        doc.add_heading(group, level=2)
        doc.add_paragraph(text["intro"])
        doc.add_paragraph(f"Impacto: {text['impact']}")
        doc.add_paragraph(f"Como corrigir: {text['fix']}")
        bullet(doc, f"Produtos nesta parte: {len(rows)}")

        for product, items in rows:
            add_product_block(doc, product, items)

        doc.add_page_break()

    doc.add_heading("Categorias vazias no menu", level=2)
    doc.add_paragraph("Estas categorias aparecem para o cliente, mas abrem paginas sem produtos. O ideal e remover do menu enquanto estiverem vazias.")
    for name, url in EMPTY_MENU_CATEGORIES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name).bold = True
        p.add_run(" | ")
        add_hyperlink(p, "abrir categoria", url)

    doc.add_heading("Seguranca e configuracao", level=2)
    doc.add_paragraph("Isto nao e pentest. E uma lista passiva do que aparece publicamente e deveria ser revisto por quem gere o WordPress/servidor.")
    for kind, detail, expected in SECURITY_FINDINGS:
        bullet(doc, f"{kind}: {detail} | Corrigir: {expected}")

    doc.add_page_break()
    doc.add_heading("Anexo visual - uma imagem por produto com erro", level=2)
    doc.add_paragraph("Cada imagem abaixo resume os erros daquele produto e mostra o link completo.")
    products_by_id = {product["id"]: product for product in products}
    for index, product_id in enumerate(sorted(cards), 1):
        product = products_by_id[product_id]
        p = doc.add_paragraph()
        p.add_run(f"{index}. {clean(product.get('name'))}").bold = True
        p.add_run(" | ")
        add_hyperlink(p, "abrir produto", product["permalink"])
        doc.add_picture(str(cards[product_id]), width=Inches(6.75))

    doc.add_page_break()
    doc.add_heading("Anexo visual - menu e seguranca", level=2)
    for card, title_text, link in general_cards:
        p = doc.add_paragraph()
        p.add_run(title_text).bold = True
        p.add_run(" | ")
        add_hyperlink(p, "abrir", link)
        doc.add_picture(str(card), width=Inches(6.75))

    doc.save(OUT_DOCX)


def build_md(products: list[dict], issues: dict[int, list[dict]]):
    grouped = grouped_rows(products, issues)
    lines = [
        "# Clickmed.pt - relatorio simples de correcoes",
        "",
        "Versao sem tabelas.",
        "",
        "## Resumo direto",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com erro/inconveniente: {len(issues)}",
        f"- Total de problemas nos produtos: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias no menu: {len(EMPTY_MENU_CATEGORIES)}",
        f"- Pontos de seguranca/configuracao: {len(SECURITY_FINDINGS)}",
        "",
    ]

    for group in GROUPS:
        rows = grouped.get(group, [])
        lines += [
            f"## {group}",
            "",
            GROUP_TEXT[group]["intro"],
            "",
            f"Produtos nesta parte: {len(rows)}",
            "",
        ]
        for product, items in rows:
            lines.append(f"- [{clean(product.get('name'))}]({product['permalink']})")
            lines.append(f"  - Tags: {', '.join(product_tags(product)) or 'sem tag'}")
            lines.append(f"  - Marcas: {', '.join(product_brands(product)) or 'sem marca'}")
            lines.append(f"  - Categorias: {', '.join(product_categories(product)) or 'sem categoria'}")
            for item in items:
                line = f"  - {item['kind']}: {item['detail']}"
                if item.get("expected"):
                    line += f" | Corrigir: {item['expected']}"
                lines.append(line)
            lines.append("")

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_products()
    issues = collect_product_issues(products)
    build_doc(products, issues)
    build_md(products, issues)
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
