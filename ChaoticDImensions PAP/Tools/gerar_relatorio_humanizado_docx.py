from __future__ import annotations

import csv
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gerar_relatorio_completo_docx import (
    CARD_DIR,
    EMPTY_MENU_CATEGORIES,
    OUT_CSV,
    SECURITY_FINDINGS,
    add_hyperlink,
    clean,
    collect_product_issues,
    load_products,
    make_non_product_card,
    make_product_card,
    product_brands,
    product_categories,
    product_tags,
)


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "relatorio_clickmed_humanizado.docx"
OUT_MD = ROOT / "relatorio_clickmed_humanizado.md"
TYPE_DIR = ROOT / "listas_clickmed_por_tipo"


GROUPS = {
    "Estado, tags e categorias": {
        "Tag/Grade",
        "Tag/condicao",
        "Categoria",
        "Descricao/Grade",
        "Tags",
        "Categorias",
    },
    "Marcas e filtros": {"Marca"},
    "Descricoes dos produtos": {
        "Descricao",
        "Descricao/IA",
        "Descricao/avaliacao",
        "Descricao/especificacao",
        "Descricao/campos vazios",
        "Descricao/imagem externa",
    },
    "Imagens e URLs": {"Imagem", "URL/slug"},
    "Organizacao da loja": {"SKU", "Promocao antiga", "Padronizacao", "Duplicado"},
}


GROUP_TEXT = {
    "Estado, tags e categorias": {
        "intro": "Aqui entram os erros que confundem o cliente antes mesmo de abrir a descricao: usado marcado como novo, Grade A/B/C sem tag correta, produto usado dentro de categoria de novo, ou produto sem categoria/tag.",
        "impact": "Isto afeta filtros, pesquisa interna, confianca e atendimento. O cliente pode comprar achando que o estado e outro.",
        "fix": "Padronizar as etiquetas para Novo, Open Box, Usado Grade A, Usado Grade B e Usado Grade C. Depois rever categorias de novo/usado.",
    },
    "Marcas e filtros": {
        "intro": "Aqui ficam os casos em que o produto aparece com marca errada ou sem marca. Exemplo claro: Samsung atribuido a Apple, Xiaomi atribuido a Sony, Nintendo atribuido a Sony.",
        "impact": "O produto aparece no filtro errado, pode mostrar icone errado e passa uma imagem de loja pouco cuidada.",
        "fix": "Corrigir a marca do produto no WooCommerce e rever os filtros/brand swatches.",
    },
    "Descricoes dos produtos": {
        "intro": "As descricoes tem muitos sinais de template/copypaste: avaliacao 4.8/5 sem reviews reais, restos de IA, campos vazios e especificacoes erradas.",
        "impact": "Esta e a parte que mais prejudica confianca. Um cliente atento percebe rapidamente quando o texto parece inventado.",
        "fix": "Remover avaliacoes falsas, apagar restos de IA, rever especificacoes e deixar textos mais curtos e factuais.",
    },
    "Imagens e URLs": {
        "intro": "Aqui entram imagens/alt/nomes de ficheiro com capacidade ou modelo diferente do produto, e URLs que dizem uma coisa enquanto o titulo diz outra.",
        "impact": "Pode gerar reclamação: o cliente ve 512GB numa imagem, mas compra um produto de 256GB, ou entra numa URL que parece outro produto.",
        "fix": "Corrigir imagens, alt text e slugs. Para produtos ja indexados, criar redirecionamento 301 quando mudar URL.",
    },
    "Organizacao da loja": {
        "intro": "Aqui estao problemas que nao impedem a compra, mas deixam a loja menos profissional: SKU ausente, produtos duplicados, promocoes antigas e mistura de idiomas nos titulos.",
        "impact": "Atrasa gestao interna, complica stock e deixa a loja com aspeto desorganizado.",
        "fix": "Criar regra de nome, SKU, campanha e duplicados. Depois aplicar em lote.",
    },
}


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.color.rgb = RGBColor(35, 35, 35)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def group_for_kind(kind: str) -> str:
    for group, kinds in GROUPS.items():
        if kind in kinds:
            return group
    return "Outros"


def grouped_rows(products: list[dict], issues: dict[int, list[dict]]):
    products_by_id = {product["id"]: product for product in products}
    rows = defaultdict(list)
    for product_id, items in issues.items():
        product = products_by_id[product_id]
        by_group = defaultdict(list)
        for item in items:
            by_group[group_for_kind(item["kind"])].append(item)
        for group, group_items in by_group.items():
            rows[group].append((product, group_items))
    for group in rows:
        rows[group].sort(key=lambda row: clean(row[0].get("name")).lower())
    return rows


def issue_summary(items: list[dict]) -> str:
    lines = []
    for item in items:
        line = f"{item['kind']}: {item['detail']}"
        if item.get("expected"):
            line += f" | Corrigir: {item['expected']}"
        lines.append(line)
    return "\n".join(lines)


def add_small_table(doc: Document, headers: list[str], rows: list[list[str | tuple[str, str]]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        run = table.rows[0].cells[index].paragraphs[0].add_run(header)
        run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            if isinstance(value, tuple):
                text, url = value
                add_hyperlink(cells[index].paragraphs[0], text, url)
            else:
                cells[index].text = value
    return table


def write_type_csvs(group_rows):
    TYPE_DIR.mkdir(exist_ok=True)
    for file in TYPE_DIR.glob("*.csv"):
        file.unlink()
    for group, rows in group_rows.items():
        safe = (
            group.lower()
            .replace(" ", "_")
            .replace(",", "")
            .replace("ç", "c")
            .replace("ã", "a")
            .replace("õ", "o")
        )
        path = TYPE_DIR / f"{safe}.csv"
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle, delimiter=";")
            writer.writerow(["id", "produto", "link", "tags", "marcas", "categorias", "problemas"])
            for product, items in rows:
                writer.writerow(
                    [
                        product["id"],
                        clean(product.get("name")),
                        product["permalink"],
                        ", ".join(product_tags(product)),
                        ", ".join(product_brands(product)),
                        ", ".join(product_categories(product)),
                        " | ".join(f"{item['kind']}: {item['detail']}" for item in items),
                    ]
                )


def ensure_cards(products: list[dict], issues: dict[int, list[dict]]) -> dict[int, Path]:
    products_by_id = {product["id"]: product for product in products}
    cards = {}
    for product_id in sorted(issues):
        cards[product_id] = make_product_card(products_by_id[product_id], issues[product_id])
    return cards


def general_cards():
    cards = []
    index = 1
    for name, url in EMPTY_MENU_CATEGORIES:
        card = make_non_product_card(
            name,
            url,
            "Categoria aparece no menu, mas esta vazia/sem produtos.",
            "Remover do menu ou esconder ate ter produtos.",
            index,
        )
        cards.append((card, f"Categoria vazia: {name}", url))
        index += 1
    for kind, detail, expected in SECURITY_FINDINGS:
        card = make_non_product_card(f"Seguranca - {kind}", "https://clickmed.pt/", detail, expected, index)
        cards.append((card, f"Seguranca: {detail}", "https://clickmed.pt/"))
        index += 1
    return cards


def build_doc(products: list[dict], issues: dict[int, list[dict]]):
    group_rows = grouped_rows(products, issues)
    write_type_csvs(group_rows)
    cards = ensure_cards(products, issues)
    cards_general = general_cards()

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Revisao Clickmed.pt - lista de correcoes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Data: 02/07/2026. Levantamento publico do site, loja, produtos, categorias e headers.").italic = True

    doc.add_heading("Antes de corrigir", level=2)
    doc.add_paragraph(
        "Este documento foi reorganizado para ser usado como checklist de trabalho. "
        "A primeira parte separa os problemas por area. No final existe um anexo visual com uma imagem para cada produto com erro."
    )
    doc.add_paragraph(
        "Sugestao pratica: corrigir primeiro estado/tags, marcas e descricoes. Depois limpar URLs, imagens, categorias vazias e organizacao interna."
    )

    total_issues = sum(len(items) for items in issues.values())
    kind_counter = Counter(item["kind"] for items in issues.values() for item in items)
    doc.add_heading("Resumo rapido", level=2)
    add_small_table(
        doc,
        ["Ponto", "Resultado"],
        [
            ["Produtos analisados", str(len(products))],
            ["Produtos com erro/inconveniente", str(len(issues))],
            ["Total de problemas encontrados nos produtos", str(total_issues)],
            ["Categorias vazias no menu", str(len(EMPTY_MENU_CATEGORIES))],
            ["Pontos de seguranca listados", str(len(SECURITY_FINDINGS))],
            ["Arquivo CSV completo", str(OUT_CSV.name)],
        ],
    )

    doc.add_heading("Contagem por tipo", level=2)
    add_small_table(doc, ["Tipo", "Quantidade"], [[kind, str(count)] for kind, count in kind_counter.most_common()])

    doc.add_page_break()

    for group in GROUPS:
        rows = group_rows.get(group, [])
        text = GROUP_TEXT[group]
        doc.add_heading(group, level=2)
        doc.add_paragraph(text["intro"])
        doc.add_paragraph(f"Impacto: {text['impact']}")
        doc.add_paragraph(f"Como corrigir: {text['fix']}")
        doc.add_paragraph(f"Produtos nesta secao: {len(rows)}")

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["ID", "Produto", "Problema", "Corrigir", "Link"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].paragraphs[0].add_run(header).bold = True
        for product, items in rows:
            cells = table.add_row().cells
            cells[0].text = str(product["id"])
            cells[1].text = clean(product.get("name"))
            cells[2].text = "\n".join(f"{item['kind']}: {item['detail']}" for item in items)
            cells[3].text = "\n".join(item.get("expected") or "Rever manualmente" for item in items)
            add_hyperlink(cells[4].paragraphs[0], "Abrir", product["permalink"])

        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Categorias vazias no menu", level=2)
    doc.add_paragraph("Estas paginas existem e aparecem no menu, mas nao mostram produtos. O ideal e remover do menu ate haver stock.")
    add_small_table(doc, ["Categoria", "Link"], [[name, ("Abrir", url)] for name, url in EMPTY_MENU_CATEGORIES])

    doc.add_heading("Seguranca", level=2)
    doc.add_paragraph(
        "Esta parte e uma revisao passiva. Nao foi pentest. Sao pontos de configuracao e exposicao que devem ser revistos com quem gere o WordPress/servidor."
    )
    add_small_table(doc, ["Area", "Problema", "Recomendacao"], [[kind, detail, expected] for kind, detail, expected in SECURITY_FINDINGS])

    doc.add_page_break()
    doc.add_heading("Anexo visual - produtos com erro", level=2)
    doc.add_paragraph(
        "Abaixo fica uma evidencia por produto. A tabela por area e melhor para corrigir; este anexo serve para confirmar rapidamente cada caso."
    )
    for index, product_id in enumerate(sorted(cards), 1):
        product = next(product for product in products if product["id"] == product_id)
        par = doc.add_paragraph()
        par.add_run(f"{index}. {clean(product.get('name'))}").bold = True
        par.add_run(" - ")
        add_hyperlink(par, "abrir produto", product["permalink"])
        doc.add_picture(str(cards[product_id]), width=Inches(6.75))

    doc.add_page_break()
    doc.add_heading("Anexo visual - menu e seguranca", level=2)
    for card, title_text, link in cards_general:
        par = doc.add_paragraph()
        par.add_run(title_text).bold = True
        par.add_run(" - ")
        add_hyperlink(par, "abrir", link)
        doc.add_picture(str(card), width=Inches(6.75))

    doc.save(OUT_DOCX)


def build_md(products: list[dict], issues: dict[int, list[dict]]):
    group_rows = grouped_rows(products, issues)
    lines = [
        "# Revisao Clickmed.pt - lista de correcoes",
        "",
        "Data: 02/07/2026.",
        "",
        "Este ficheiro e uma versao mais limpa do relatorio. Os detalhes completos tambem estao no DOCX e nos CSVs separados por tipo.",
        "",
        "## Resumo",
        "",
        f"- Produtos analisados: {len(products)}",
        f"- Produtos com erro/inconveniente: {len(issues)}",
        f"- Total de problemas nos produtos: {sum(len(items) for items in issues.values())}",
        f"- Categorias vazias no menu: {len(EMPTY_MENU_CATEGORIES)}",
        "",
    ]
    for group in GROUPS:
        rows = group_rows.get(group, [])
        lines += [f"## {group}", "", GROUP_TEXT[group]["intro"], "", f"Produtos nesta secao: {len(rows)}", ""]
        for product, items in rows:
            lines.append(f"- [{clean(product.get('name'))}]({product['permalink']})")
            for item in items:
                line = f"  - {item['kind']}: {item['detail']}"
                if item.get("expected"):
                    line += f" | Corrigir: {item['expected']}"
                lines.append(line)
        lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    products = load_products()
    issues = collect_product_issues(products)
    build_doc(products, issues)
    build_md(products, issues)
    print(OUT_DOCX)
    print(OUT_MD)
    print(TYPE_DIR)


if __name__ == "__main__":
    main()
