from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "relatorio_clickmed.md"
OUT_PATH = ROOT / "relatorio_clickmed_com_imagens.docx"
SHOT_DIR = ROOT / "screenshots_clickmed"
ANN_DIR = SHOT_DIR / "annotated"


def font(size=28, bold=False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font_obj, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = (current + " " + word).strip()
        width = draw.textbbox((0, 0), test, font=font_obj)[2]
        if width <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def annotate_image(source, dest, note, boxes=None, crop=None):
    boxes = boxes or []
    image = Image.open(source).convert("RGB")
    offset_x = offset_y = 0
    if crop:
        left, top, right, bottom = crop
        image = image.crop(crop)
        offset_x = left
        offset_y = top

    banner_h = 120
    out = Image.new("RGB", (image.width, image.height + banner_h), "white")
    out.paste(image, (0, banner_h))

    draw = ImageDraw.Draw(out)
    title_font = font(31, bold=True)
    small_font = font(24)
    draw.rectangle((0, 0, out.width, banner_h), fill=(255, 246, 240))
    draw.rectangle((0, banner_h - 4, out.width, banner_h), fill=(255, 105, 0))

    lines = wrap_text(draw, note, title_font, out.width - 60)
    y = 18
    for i, line in enumerate(lines[:2]):
        draw.text((30, y), line, fill=(170, 35, 0), font=title_font if i == 0 else small_font)
        y += 42

    for box in boxes:
        x1, y1, x2, y2 = box
        x1 = x1 - offset_x
        x2 = x2 - offset_x
        y1 = y1 - offset_y + banner_h
        y2 = y2 - offset_y + banner_h
        draw.rectangle((x1, y1, x2, y2), outline=(220, 0, 0), width=8)

    out.save(dest, quality=92)


def prepare_images():
    ANN_DIR.mkdir(parents=True, exist_ok=True)
    specs = [
        {
            "file": "home_categories_products.png",
            "out": "01_home_categoria_smart_home.png",
            "note": "Home mostra Smart Home nas categorias, mas a categoria esta vazia.",
            "crop": (0, 0, 1365, 1180),
            "boxes": [(650, 875, 780, 1045), (80, 1180, 320, 1585)],
            "caption": "Home: categoria Smart Home aparece para o cliente; depois abre sem produtos.",
        },
        {
            "file": "empty_smart_home.png",
            "out": "02_categoria_vazia.png",
            "note": "Categoria Smart Home: pagina existe, mas nao tem nenhum produto.",
            "crop": (0, 170, 1365, 680),
            "boxes": [(350, 360, 1335, 430)],
            "caption": "Categoria vazia no menu: Smart Home retorna 200, mas mostra zero produtos.",
        },
        {
            "file": "samsung_brand_apple.png",
            "out": "03_samsung_marca_apple.png",
            "note": "Produto Samsung com marca/icone Apple no cartao do produto.",
            "crop": (0, 240, 1365, 1200),
            "boxes": [(965, 260, 1080, 350), (960, 360, 1300, 430)],
            "caption": "Exemplo de marca errada: Samsung S24 Ultra aparece com Apple como marca.",
        },
        {
            "file": "iphone16_tag_novo.png",
            "out": "04_avaliacao_falsa_meses.png",
            "note": "Descricao mostra 4.8/5 e '18 meses meses' em produto sem reviews reais.",
            "crop": (0, 980, 1365, 1750),
            "boxes": [(65, 1285, 320, 1385), (65, 1660, 430, 1715)],
            "caption": "Descricao gerada/repetida: avaliacao 4.8/5 sem reviews reais e erro '18 meses meses'.",
        },
        {
            "file": "samsung_s25_a14.png",
            "out": "05_samsung_a14_bionic.png",
            "note": "Samsung Open Box com texto errado: A14 Bionic, campos vazios e meses repetidos.",
            "crop": (0, 950, 1365, 1970),
            "boxes": [(65, 1195, 320, 1290), (65, 1545, 520, 1660), (65, 1770, 1250, 1845)],
            "caption": "Erro grave de descricao: produto Samsung menciona chip A14 Bionic, que e da Apple.",
        },
        {
            "file": "samsung_s26_url_image.png",
            "out": "06_s26_url_capacidade.png",
            "note": "Produto diz 16GB/1TB, mas a URL/slug usa 12gb e os dados da imagem indicam 12GB/512GB.",
            "crop": (0, 220, 1365, 980),
            "boxes": [(960, 360, 1310, 455), (85, 270, 190, 550)],
            "caption": "Exemplo de URL/imagem/capacidade incoerente no Samsung S26 Ultra.",
        },
    ]

    ready = []
    for spec in specs:
        source = SHOT_DIR / spec["file"]
        dest = ANN_DIR / spec["out"]
        annotate_image(source, dest, spec["note"], spec["boxes"], spec["crop"])
        ready.append((dest, spec["caption"]))
    return ready


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_markdown_line(doc, line):
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        return

    numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if numbered:
        paragraph = doc.add_paragraph(style="List Number")
        add_inline_runs(paragraph, numbered.group(2))
        return

    if stripped.startswith("- "):
        paragraph = doc.add_paragraph(style="List Bullet")
        add_inline_runs(paragraph, stripped[2:])
        return

    paragraph = doc.add_paragraph()
    add_inline_runs(paragraph, stripped)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def style_document(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(30, 30, 30)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def build_docx(images):
    doc = Document()
    style_document(doc)

    title = doc.add_heading("Relatorio Clickmed.pt com imagens", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Gerado em 02/07/2026. Analise publica/passiva; sem pentest invasivo.").italic = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Resumo"
    hdr[1].text = "Numero"
    set_cell_shading(hdr[0], "F2F2F2")
    set_cell_shading(hdr[1], "F2F2F2")
    rows = [
        ("Produtos analisados pela API publica", "417"),
        ("Produtos com Grade no titulo sem tag exata", "119"),
        ("Conflitos titulo/tag/categoria", "17 + 19"),
        ("Marcas provavelmente erradas", "19"),
        ("Imagens/alt com modelo ou capacidade diferente", "29"),
        ("Descricoes com 4.8/5 apesar de 0 reviews reais", "238"),
        ("Categorias vazias no menu", "14"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right

    doc.add_heading("Imagens dos problemas", level=2)
    for image_path, caption in images:
        cap = doc.add_paragraph()
        cap.add_run(caption).bold = True
        doc.add_picture(str(image_path), width=Inches(6.9))

    doc.add_page_break()
    doc.add_heading("Relatorio textual", level=2)
    for line in MD_PATH.read_text(encoding="utf-8").splitlines():
        add_markdown_line(doc, line)

    doc.save(OUT_PATH)


def main():
    images = prepare_images()
    build_docx(images)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
