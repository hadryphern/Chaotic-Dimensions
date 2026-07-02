#!/usr/bin/env python3
"""Register the 245-item expansion and its concept boards in the PAP report."""

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "ChaoticDImensionsMod_Relatório_Pap.docx"
BACKUPS = ROOT / "tmp" / "report_backups"
ART = ROOT / "assets_work" / "concept_sketches"
MARKER = "Catálogo funcional de 245 itens"


def replace_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the paragraph and first-run formatting."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def clear_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def add_subheading(document: Document, text: str, page_break: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.page_break_before = page_break
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return paragraph


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return paragraph


def add_picture_pair(document: Document, left: Path, right: Path) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    clear_table_borders(table)
    for cell, path in zip(table.rows[0].cells, (left, right)):
        cell.width = Inches(3.15)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(3.0))


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def insert_development_section(document: Document) -> None:
    target = next(p for p in document.paragraphs if p.text.strip() == "Conteúdo legacy e mobs")
    heading = target.insert_paragraph_before(MARKER, style="Heading 2")
    heading.paragraph_format.keep_with_next = True

    target.insert_paragraph_before(
        "Foi implementado um catálogo funcional de 245 itens distribuídos por oito famílias: "
        "50 itens melee, 45 ranged, 45 magic, 45 summon, 25 acessórios, 15 ferramentas, "
        "10 consumíveis e 10 materiais. O conjunto acompanha 17 etapas, desde o início do "
        "jogo até ao conteúdo posterior ao Alien Kraken.",
        style="PAP Body",
    )
    target.insert_paragraph_before(
        "As armas têm dano, velocidade, custo de mana, munição, alcance e comportamento ajustados "
        "ao respetivo patamar. Os itens de summon incluem minions, chicotes e beacons; os ranged "
        "usam arcos, carabinas e lançadores; melee e magic alternam ataques diretos e projéteis. "
        "As receitas verificam a progressão e reutilizam materiais de Terraria e do mod.",
        style="PAP Body",
    )
    target.insert_paragraph_before(
        "Os tiers finais foram deliberadamente tratados como conteúdo de teste extremo. A linha "
        "Krakenbane, desbloqueada depois do Crystaline Devourer, alcança 12 500 000 de dano base e "
        "pode derrotar o Kraken atual num golpe; as linhas Abyssal e Chaotic sobem para 25 000 000 "
        "e 50 000 000. O Kraken passou também a deixar Abyssal Kraken Core, material necessário às "
        "receitas pós-boss.",
        style="PAP Body",
    )
    target.insert_paragraph_before(
        "Enquanto não existem sprites exclusivas, cada item utiliza uma textura vanilla coerente "
        "como placeholder. O comportamento comum ficou concentrado em classes base e em quatro "
        "projéteis partilhados, reduzindo repetição sem transformar os itens em cópias idênticas. "
        "O catálogo nominal completo encontra-se em docs/Catalogo_245_Itens.md e no Anexo C.",
        style="PAP Body",
    )


def append_concept_annex(document: Document) -> None:
    add_subheading(document, "C.3 Expansão de itens e conceitos visuais", page_break=True)
    paragraph = document.add_paragraph(style="PAP Body")
    paragraph.add_run(
        "As pranchas seguintes registam a expansão posterior à primeira auditoria visual. Reúnem "
        "os bosses atuais e planeados, uma proposta de mobs futuros, o percurso dos 17 tiers e os "
        "245 itens funcionais. As imagens são esboços de produção e não substituem as sprites finais."
    )

    add_caption(document, "Prancha C.19 - Esquema técnico dos bosses")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ART / "Boss_Roster_Blueprint.png"), width=Inches(5.8))
    source = document.add_paragraph("Fonte: elaboração própria a partir do esquema funcional do projeto.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)

    document.add_page_break()
    add_caption(document, "Prancha C.20 - Mobs atuais, planeados e progressão dos itens")
    add_picture_pair(document, ART / "Mob_Roster_Blueprint.png", ART / "Item_Progression_Blueprint.png")
    source = document.add_paragraph("Fonte: elaboração própria a partir do roteiro de progressão.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)

    document.add_page_break()
    add_caption(document, "Prancha C.21 - Catálogo dos itens 1 a 200")
    add_picture_pair(document, ART / "Item_Sheets_Overview_A.jpg", ART / "Item_Sheets_Overview_B.jpg")
    source = document.add_paragraph("Fonte: catálogo gerado a partir de Content/Items/Progression.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)

    document.add_page_break()
    add_caption(document, "Prancha C.22 - Catálogo dos itens 201 a 245")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ART / "Item_Sheets_Overview_C.jpg"), width=Inches(5.8))
    source = document.add_paragraph("Fonte: catálogo gerado a partir de Content/Items/Progression.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)


def main() -> None:
    document = Document(REPORT)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        raise SystemExit("The 245-item report section already exists; no changes made.")

    required = [
        ART / "Boss_Roster_Blueprint.png",
        ART / "Mob_Roster_Blueprint.png",
        ART / "Item_Progression_Blueprint.png",
        ART / "Item_Sheets_Overview_A.jpg",
        ART / "Item_Sheets_Overview_B.jpg",
        ART / "Item_Sheets_Overview_C.jpg",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise SystemExit("Missing concept boards: " + ", ".join(missing))

    BACKUPS.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(REPORT, BACKUPS / f"{REPORT.stem}_before_245_items_{stamp}.docx")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("O Chaotic Dimensions Mod expande Terraria"):
            replace_text(
                paragraph,
                "O Chaotic Dimensions Mod expande Terraria com uma identidade própria baseada em "
                "dimensões caóticas, cristais, sombra, criaturas alienígenas e conteúdo legado. A "
                "versão analisada contém 131 ficheiros C#, 518 tipos declarados e um novo catálogo "
                "funcional de 245 itens, além dos bosses, mobs, projéteis, buffs, tiles, sistemas e "
                "quatro faixas de música já integrados.",
            )
        elif text.startswith("A análise atualizada identificou 117 ficheiros C#"):
            replace_text(
                paragraph,
                "A análise atualizada identificou 131 ficheiros C# e 518 tipos entre classes, "
                "estruturas e enums. A expansão acrescenta 245 itens funcionais, receitas associadas, "
                "três projéteis de armas, um projétil de minion, um buff e o catálogo técnico usado "
                "para controlar os 17 patamares de progressão.",
            )
        elif text.startswith("O repositório contém 117 ficheiros C#"):
            replace_text(
                paragraph,
                "O repositório contém 131 ficheiros C#, abrangendo itens, projéteis, NPCs, buffs, "
                "tiles, sistemas e classes auxiliares. O código totaliza 645 linhas de comentário e "
                "é acompanhado por docs/Guiao_Apresentacao_Codigo.md, que explica os bosses, o "
                "vocabulário do tModLoader e a função de cada ficheiro.",
            )
        elif text.startswith("Este anexo reúne o material visual histórico"):
            replace_text(
                paragraph,
                "Este anexo reúne o material visual histórico e o catálogo atual de sprites. As "
                "figuras recuperadas do relatório anterior foram compactadas em grelhas, mantendo "
                "os bookmarks do Índice de Figuras. As primeiras dezoito pranchas correspondem aos "
                "322 recursos da auditoria inicial; a secção C.3 regista a expansão posterior.",
            )
        elif text.startswith("As pranchas apresentam todos os 322 ficheiros visuais"):
            replace_text(
                paragraph,
                "As pranchas C.01 a C.18 apresentam os 322 ficheiros visuais da primeira auditoria. "
                "Cada célula identifica o nome, a pasta e a resolução original; spritesheets e GIFs "
                "são representados por um frame de referência.",
            )

    insert_development_section(document)
    append_concept_annex(document)
    set_update_fields(document)
    document.save(REPORT)
    print(f"Updated report: {REPORT}")


if __name__ == "__main__":
    main()
