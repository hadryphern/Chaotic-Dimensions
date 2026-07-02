# -*- coding: latin-1 -*-
from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "ChaoticDImensionsMod_Relatório_Pap.docx"
SHEETS = ROOT / "assets_work" / "report_annex_sheets"
BACKUP = ROOT / "tmp" / "report_backups" / f"ChaoticDImensionsMod_Relatório_Pap_before_compact_annex_{datetime.now():%Y%m%d_%H%M%S}.docx"
BACKUP.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(REPORT, BACKUP)

document = Document(REPORT)
paragraphs = document.paragraphs
if len(paragraphs) < 600 or paragraphs[419].text.strip() != "Anexo C - Registo visual do relatório anterior":
    raise RuntimeError("Unexpected annex structure")

# Keep all existing figure bookmarks and image relationships by moving the original XML nodes.
records = []
for start in range(421, 600, 4):
    caption = paragraphs[start]._p
    picture = paragraphs[start + 1]._p
    source = paragraphs[start + 2]._p
    records.append((caption, picture, source))
if len(records) != 45:
    raise RuntimeError(f"Expected 45 recovered figures, found {len(records)}")

# Rename the existing Annex C entry without disturbing its bookmarks/hyperlinks.
def replace_text_nodes(paragraph, old, new):
    texts = paragraph._p.xpath('.//*[local-name()="t"]')
    combined = "".join(node.text or "" for node in texts)
    if old not in combined:
        raise RuntimeError(f"Text not found: {old}")
    if len(texts) == 1:
        texts[0].text = combined.replace(old, new, 1)
        return
    replacement = combined.replace(old, new, 1)
    texts[0].text = replacement
    for node in texts[1:]:
        node.text = ""

replace_text_nodes(paragraphs[114], "11.3 Anexo C - Registo visual do relatório anterior", "11.3 Anexo C - Registo visual e catálogo de sprites")
replace_text_nodes(paragraphs[419], "Anexo C - Registo visual do relatório anterior", "Anexo C - Registo visual e catálogo de sprites")
paragraphs[420].text = (
    "Este anexo reúne o material visual histórico e o catálogo atual de sprites. "
    "As figuras recuperadas do relatório anterior foram compactadas em grelhas, mantendo os bookmarks do Índice de Figuras. "
    "As pranchas seguintes inventariam 322 recursos visuais existentes em Content, Assets e assets_work, incluindo sprites integradas, fontes e versões de evolução."
)

# Remove the old one-image-per-page sequence from the document body.
body = document._element.body
for paragraph in paragraphs[421:]:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)

def add_page_break():
    p = document.add_paragraph()
    p.add_run().add_break()
    # Convert the default line break to a page break.
    br = p._p.xpath('.//*[local-name()="br"]')[0]
    br.set(qn("w:type"), "page")
    return p

def set_cell_margins(cell, top=45, start=45, bottom=45, end=45):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_paragraph_font(paragraph_element, size_half_points, italic=False):
    for run in paragraph_element.xpath('.//*[local-name()="r"]'):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run.insert(0, rPr)
        for tag in ("w:sz", "w:szCs"):
            element = rPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                rPr.append(element)
            element.set(qn("w:val"), str(size_half_points))
        fonts = rPr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rPr.append(fonts)
        for attr in ("ascii", "hAnsi", "eastAsia"):
            fonts.set(qn(f"w:{attr}"), "Arial")
        if italic and rPr.find(qn("w:i")) is None:
            rPr.append(OxmlElement("w:i"))

def set_spacing(paragraph_element, before=0, after=0, line=180):
    pPr = paragraph_element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    justification = pPr.find(qn("w:jc"))
    if justification is None:
        justification = OxmlElement("w:jc")
        pPr.append(justification)
    justification.set(qn("w:val"), "center")

def resize_drawing(paragraph_element, max_width, max_height):
    extents = paragraph_element.xpath('.//*[local-name()="extent"]')
    transforms = paragraph_element.xpath('.//*[local-name()="xfrm"]/*[local-name()="ext"]')
    if not extents:
        return
    cx = int(extents[0].get("cx"))
    cy = int(extents[0].get("cy"))
    scale = min(max_width / cx, max_height / cy, 1.0)
    new_cx = max(1, int(cx * scale))
    new_cy = max(1, int(cy * scale))
    for extent in extents:
        extent.set("cx", str(new_cx))
        extent.set("cy", str(new_cy))
    for extent in transforms:
        extent.set("cx", str(new_cx))
        extent.set("cy", str(new_cy))

add_page_break()
section_title = document.add_paragraph("C.1 Registo visual recuperado")
section_title.style = "Normal"
for run in section_title.runs:
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)
intro = document.add_paragraph(
    "As 45 figuras seguintes mantêm a numeração e as ligações do Índice de Figuras, mas passam a ocupar nove posições por prancha."
)
intro.style = "PAP Body"

for table_index, offset in enumerate(range(0, len(records), 9)):
    if table_index > 0:
        add_page_break()
    batch = records[offset:offset + 9]
    table = document.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D5DDE7")
        borders.append(border)
    tblPr.append(borders)
    for row in table.rows:
        row.height = Inches(2.38)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)

    for index in range(9):
        row, col = divmod(index, 3)
        cell = table.cell(row, col)
        cell.width = Inches(2.05)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:p"):
                tc.remove(child)

        if index >= len(batch):
            tc.append(OxmlElement("w:p"))
            continue

        caption, picture, source = batch[index]
        resize_drawing(picture, int(Inches(1.82)), int(Inches(1.32)))
        set_paragraph_font(caption, 14)
        set_paragraph_font(source, 11, italic=True)
        set_spacing(caption, after=10, line=160)
        set_spacing(picture, after=10, line=160)
        set_spacing(source, after=0, line=140)
        tc.append(caption)
        tc.append(picture)
        tc.append(source)

add_page_break()
catalog_title = document.add_paragraph("C.2 Catálogo atual de sprites e recursos visuais")
catalog_title.style = "Normal"
for run in catalog_title.runs:
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)
catalog_intro = document.add_paragraph(
    "As pranchas apresentam todos os 322 ficheiros visuais inventariados. "
    "Cada célula identifica o nome, a pasta e a resolução original; spritesheets e GIFs são representados por um frame de referência."
)
catalog_intro.style = "PAP Body"

sheet_paths = sorted(SHEETS.glob("sheet_*.jpg"))
if len(sheet_paths) != 18:
    raise RuntimeError(f"Expected 18 contact sheets, found {len(sheet_paths)}")

for index, path in enumerate(sheet_paths, start=1):
    if index > 1:
        add_page_break()
    title = document.add_paragraph(f"Prancha C.{index:02d} - Catálogo visual do projeto")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(path), width=Inches(6.18))

    source = document.add_paragraph(
        f"Fonte: elaboração própria a partir do inventário do projeto; ficheiro {path.name}."
    )
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in source.runs:
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(7)

# Request field refresh when the document is opened in OnlyOffice/Word.
settings = document.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

document.save(REPORT)
print(f"Saved compact annex with 45 linked figures and {len(sheet_paths)} sprite sheets")
print(f"Backup: {BACKUP}")
