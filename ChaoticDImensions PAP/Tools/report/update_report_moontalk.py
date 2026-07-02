#!/usr/bin/env python3
"""Add the MoonTalk narrative sequence to the PAP report without rebuilding it."""

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "ChaoticDImensionsMod_Relatório_Pap.docx"
PREVIEW = ROOT / "assets_work/moontalk/MoonTalk_Preview_Quatro_Estados.png"
SOUL_PREVIEW = ROOT / "assets_work/soul_orb/SoulOrb_64_Frames.png"
BACKUPS = ROOT / "tmp/report_backups"
MARKER = "MoonTalk e introdução narrativa"


def move_before(paragraph, nodes) -> None:
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for node in nodes:
        parent.insert(index, node)
        index += 1


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    if not PREVIEW.exists():
        raise SystemExit(f"Missing MoonTalk preview: {PREVIEW}")

    document = Document(REPORT)
    if any(paragraph.text.strip() == MARKER for paragraph in document.paragraphs):
        replacements = {
            "MoonTalk foi concebido como uma presença central na narrativa do mod.": (
                "A introdução narrativa apresentada após a criação de um mundo novo deixou de revelar diretamente o MoonTalk. A voz passa a ser representada por uma Alma Guia branca, mantendo a identidade do interlocutor ambígua e preservando o MoonTalk para uma futura batalha de boss. Mundos já existentes não iniciam a sequência. Durante os primeiros 31 segundos, o ecrã permanece completamente preto, a música distorcida começa e o jogador fica imóvel, sem poder usar itens ou receber dano."
            ),
            "Depois desse período, MoonTalk surge lentamente.": (
                "Depois desse período, a Alma Guia surge durante um fade de oito segundos. A versão final foi simplificada para um único núcleo branco com margem azul-fria e uma pequena chama superior, sem rosto, partículas, cauda inferior ou ornamentos. O atlas contém apenas quatro frames discretos de respiração; o movimento lateral e a flutuação vertical são interpolados em tempo real, evitando uma animação excessivamente detalhada. A aparência é aparentemente inofensiva, mas a voz mantém a presença inquietante do MoonTalk por trás do diálogo."
            ),
            "Depois desse período, a Alma Guia surge": (
                "Depois desse período, a Alma Guia surge durante um fade de oito segundos sobre um fundo totalmente preto. A versão atual utiliza duas esferas brancas, sem rosto, partículas, cauda, chama, arcos ou ornamentos. O centro branco transita suavemente para uma margem cinza-azulada muito clara. Os 64 frames mantêm as esferas a orbitar uma à outra, mas alteram progressivamente o eixo entre trajetórias horizontais, diagonais e verticais. O conjunto também sobe e desce lentamente enquanto dialoga com o jogador. Cada frame passou de 128 para 256 píxeis, reduzindo a pixelização causada pela ampliação no jogo."
            ),
            "A sequência é controlada por MoonTalkIntroSystem e guardada por MoonTalkWorldSystem.": (
                "A sequência é controlada por MoonTalkIntroSystem e guardada por MoonTalkWorldSystem. Durante a apresentação, efeitos e ambiente são silenciados, enquanto uma SceneEffect de prioridade máxima força a música vanilla para silêncio; apenas a faixa e as vozes do MoonTalk permanecem audíveis. No encerramento, personagem e fundo preto desaparecem num fade de cinco segundos, enquanto a música reduz gradualmente o volume até ao silêncio. Os volumes anteriores do jogo são restaurados depois da cena."
            ),
        }
        for paragraph in document.paragraphs:
            for prefix, text in replacements.items():
                if paragraph.text.startswith(prefix):
                    paragraph.text = text
                    paragraph.style = "PAP Body"

        caption_index = next(
            i for i, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip() == "Figura 59 - Estados visuais do MoonTalk"
        )
        picture_paragraph = document.paragraphs[caption_index + 1]
        blips = picture_paragraph._p.xpath('.//*[local-name()="blip"]')
        if blips:
            relationship_id = blips[0].get(qn("r:embed"))
            document.part.related_parts[relationship_id]._blob = PREVIEW.read_bytes()

        soul_caption = "Figura 60 - Frames da Alma Guia usada na introdução"
        if not any(paragraph.text.strip() == soul_caption for paragraph in document.paragraphs):
            tests_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Testes")
            soul_heading = document.add_paragraph("Animação da Alma Guia", style="Heading 2")
            soul_text = document.add_paragraph(
                "A prancha apresenta os oito frames base. O atlas final encontra-se em Assets/SoulOrb/SoulOrb_Atlas.png, enquanto o MoonTalk permanece separado em Assets/MoonTalk para desenvolvimento futuro como boss.",
                style="PAP Body",
            )
            caption = document.add_paragraph(soul_caption, style="Normal")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].bold = True
            caption.runs[0].font.size = Pt(10)
            picture = document.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.add_run().add_picture(str(SOUL_PREVIEW), width=Inches(5.65))
            source = document.add_paragraph("Fonte: elaboração própria e processamento da animação para tModLoader.")
            source.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source.runs[0].italic = True
            source.runs[0].font.size = Pt(9)
            move_before(tests_heading, [soul_heading._p, soul_text._p, caption._p, picture._p, source._p])
        else:
            soul_caption_index = next(
                i for i, paragraph in enumerate(document.paragraphs)
                if paragraph.text.strip() == soul_caption
            )
            for paragraph in document.paragraphs:
                if paragraph.text.startswith("A prancha apresenta os ") and "frames base" in paragraph.text:
                    paragraph.text = (
                        "A prancha apresenta os 64 frames em alta resolução das duas esferas durante a órbita multidirecional. O atlas final encontra-se em Assets/SoulOrb/SoulOrb_Atlas.png, enquanto o MoonTalk permanece separado em Assets/MoonTalk para desenvolvimento futuro como boss."
                    )
                    paragraph.style = "PAP Body"
            soul_picture = document.paragraphs[soul_caption_index + 1]
            soul_blips = soul_picture._p.xpath('.//*[local-name()="blip"]')
            if soul_blips:
                soul_relationship_id = soul_blips[0].get(qn("r:embed"))
                document.part.related_parts[soul_relationship_id]._blob = SOUL_PREVIEW.read_bytes()

        build_error_marker = "Durante um teste de Construir + Recarregar"
        if not any(paragraph.text.startswith(build_error_marker) for paragraph in document.paragraphs):
            tests_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Testes")
            correction = document.add_paragraph(
                build_error_marker + ", o compilador apresentou 6076 erros em cascata. A análise revelou que 136 ficheiros C# tinham namespaces contaminados por caminhos de tmp/comment_backups, incluindo segmentos numéricos inválidos. Foi criada uma cópia de segurança, os nomes qualificados foram reparados mecanicamente e a ferramenta de comentários passou a ignorar tmp e tools. A regra automática de correspondência entre namespace e pasta também foi desativada. A compilação integral do tModLoader terminou depois com zero erros e zero avisos.",
                style="PAP Body",
            )
            move_before(tests_heading, [correction._p])

        set_update_fields(document)
        BACKUPS.mkdir(parents=True, exist_ok=True)
        backup = BACKUPS / f"{REPORT.stem}_before_moontalk_revision_{datetime.now():%Y%m%d_%H%M%S}.docx"
        shutil.copy2(REPORT, backup)
        document.save(REPORT)
        print(f"Updated MoonTalk revision: {REPORT}")
        print(f"Backup: {backup}")
        return

    target = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Testes")
    created = []

    heading = document.add_paragraph(MARKER, style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    created.append(heading._p)

    texts = (
        "MoonTalk foi concebido como uma presença central na narrativa do mod. A sua primeira aparição acontece uma única vez em cada mundo: durante os primeiros 31 segundos, o ecrã permanece completamente preto, a música distorcida começa e o jogador fica imóvel, sem poder usar itens ou receber dano.",
        "Depois desse período, MoonTalk surge lentamente. A imagem utiliza uma sombra deslocada e um contorno luminoso muito fraco gerado a partir da própria silhueta, sem aplicar aura ao personagem. As dez falas usam gravações WAV próprias, iniciadas juntamente com cada legenda. O texto aparece letra a letra com a fonte do Terraria e permanece visível durante pelo menos cinco segundos; as falas mais longas conservam a legenda até o áudio terminar.",
        "A sequência é controlada por MoonTalkIntroSystem e guardada por MoonTalkWorldSystem. Ao concluir, o estado é gravado no mundo e sincronizado em multiplayer. Se o jogador sair antes do fim, a introdução volta a ser apresentada. A música de fundo é reproduzida com volume reduzido, as vozes foram atenuadas para evitar saturação e todos os sons são interrompidos quando a cena termina ou é cancelada.",
        "A construção visual passou por correções de encaixe. A primeira interpretação colocou os braços voltados para cima; os conceitos posteriores esclareceram que o braço normal desce a partir do peito. O segundo sprite passou a representar o braço de ataque completo. Foram então preparados quatro estados independentes: repouso, braço esquerdo levantado, braço direito levantado e ambos levantados.",
    )
    for text in texts:
        paragraph = document.add_paragraph(text, style="PAP Body")
        created.append(paragraph._p)

    caption = document.add_paragraph("Figura 59 - Estados visuais do MoonTalk", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].bold = True
    caption.runs[0].font.size = Pt(10)
    created.append(caption._p)

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(PREVIEW), width=Inches(5.65))
    created.append(picture._p)

    source = document.add_paragraph("Fonte: elaboração própria a partir das sprites do MoonTalk.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)
    created.append(source._p)

    move_before(target, created)
    set_update_fields(document)

    BACKUPS.mkdir(parents=True, exist_ok=True)
    backup = BACKUPS / f"{REPORT.stem}_before_moontalk_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(REPORT, backup)
    document.save(REPORT)
    print(f"Updated: {REPORT}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
