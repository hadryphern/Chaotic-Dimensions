# -*- coding: utf-8 -*-
"""
Gerador do Relatório PAP — Chaotic Dimensions
Corre este script para reconstruir o .docx com formatação actualizada.
Uso: python gerar_relatorio.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Relatório_PAP_Chaotic_Dimensions.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_spacing(para, before_pt=0, after_pt=12, line_rule="auto", line_val=360):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"),     str(int(after_pt  * 20)))
    spacing.set(qn("w:before"),    str(int(before_pt * 20)))
    spacing.set(qn("w:line"),      str(line_val))
    spacing.set(qn("w:lineRule"),  line_rule)
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

def apply_normal(para):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.first_line_indent = Cm(0.6)
    set_spacing(para, before_pt=0, after_pt=12, line_val=360)

def apply_h1(para, doc):
    para.style = doc.styles["Heading 1"]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.first_line_indent = Cm(0)
    set_spacing(para, before_pt=12, after_pt=12, line_val=360)
    for run in para.runs:
        run.font.name  = "Calibri Light"
        run.font.size  = Pt(16)
        run.font.bold  = False
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def apply_h2(para, doc):
    para.style = doc.styles["Heading 2"]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.first_line_indent = Cm(0)
    set_spacing(para, before_pt=6, after_pt=12, line_val=360)
    for run in para.runs:
        run.font.name  = "Calibri Light"
        run.font.size  = Pt(12)
        run.font.bold  = False
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h1(doc, text):
    p = doc.add_paragraph(text)
    apply_h1(p, doc)
    return p

def h2(doc, text):
    p = doc.add_paragraph(text)
    apply_h2(p, doc)
    return p

def np(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    apply_normal(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return p

# ── documento ────────────────────────────────────────────────────────────────

doc = Document()

# Margens: 2.5 cm cima/baixo, 3 cm esq/dir
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# Estilo Normal base
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
nf = normal.paragraph_format
nf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
nf.first_line_indent = Cm(0.6)
nf.space_after       = Pt(12)

# ─── TÍTULO ────────────────────────────────────────────────────────────────
h1(doc, "Relatório de Desenvolvimento — Chaotic Dimensions")

np(doc, (
    "Documento de acompanhamento do desenvolvimento do mod Chaotic Dimensions, "
    "elaborado no âmbito da Prova de Aptidão Profissional (PAP). "
    "Este relatório regista, de forma cronológica, todas as alterações, adições, "
    "remoções, erros e decisões técnicas tomadas ao longo do projeto."
))

# ─── 1. INTRODUÇÃO ─────────────────────────────────────────────────────────
h1(doc, "1. Introdução ao Projeto")

np(doc, (
    "O Chaotic Dimensions é um mod original para o jogo Terraria, desenvolvido "
    "através da plataforma tModLoader. O projeto foi concebido como um universo "
    "próprio, distinto de simples pacotes de conteúdo, incorporando bosses "
    "originais, biomas personalizados e sistemas de progressão inéditos."
))
np(doc, (
    "O desenvolvimento é realizado em C# e constitui o trabalho central da Prova "
    "de Aptidão Profissional. O mod encontra-se na versão 0.2 e o autor "
    "identifica-se sob o pseudónimo blueDev."
))

# ─── 2. ESTADO INICIAL ─────────────────────────────────────────────────────
h1(doc, "2. Estado Inicial do Projeto (Análise — 11 de Junho de 2026)")

h2(doc, "2.1. Bosses Implementados")
np(doc, (
    "À data do início deste registo, o projeto contava com três bosses funcionais: "
    "Monthra (pré-hardmode, 22 000 HP, dois estados de ataque — HoverVolley e "
    "SweepingBurst), Crystaline Devourer (pós-Moon Lord, worm boss segmentado com "
    "intro cinemática e arena personalizada) e Alien Kraken (85 000 000 HP, três "
    "fases de combate, sete estados de ataque distintos e sistema de evento completo "
    "com suporte a multijogador)."
))
np(doc, (
    "Encontravam-se ainda registados no sistema de gravação do mundo três bosses "
    "planeados — ChaoticApexOne, ChaoticApexTwo e ChaoticApexThree — sem "
    "implementação de código NPC correspondente, representando conteúdo previsto "
    "para versões futuras."
))

h2(doc, "2.2. Progressão e Conteúdo")
np(doc, (
    "A cadeia de progressão do mod seguia a seguinte ordem: Monthra (drop: "
    "MonthraScale) → Arsenal Rosalita no Shadow Biome → itens Shadow e "
    "EclipsedMonthra → Crystaline Devourer (drop: CrystalineTear, armadura "
    "CrystalineDevour) → Alien Kraken, que requeria o Crystaline Devourer derrotado. "
    "O Bioma Shadow dispunha de geração de mundo personalizada, NPCs próprios "
    "(ShadowSlime, ShadowEye, ShadowWorm, KrakenSquid, Phantasm) e uma linha "
    "completa de itens dedicada."
))

h2(doc, "2.3. Sistemas e Infraestrutura")
np(doc, (
    "O projeto incluía os seguintes sistemas principais: KrakenEventSystem para "
    "gestão completa do evento do Kraken (temporizador, intro com título, efeitos "
    "de tinta no ecrã, agitação de câmara, névoa alienígena, máscara de visão, "
    "grau de combate, transição de fase 2 e sincronização em rede); "
    "ChaoticDownedBossSystem para persistência de flags de bosses derrotados por "
    "mundo; sistemas de intro e arena para Monthra e Crystaline Devourer; sistema "
    "de contagem e geração do Bioma Shadow; e ChaoticProgressionGate para controlo "
    "de progressão."
))

h2(doc, "2.4. Conteúdo Incompleto e Placeholders")
np(doc, (
    "A intro cinemática do Alien Kraken encontrava-se desativada, com as constantes "
    "de cutscene definidas como 9999 × 60 no KrakenEventSystem. O conteúdo "
    "MinecraftLegacy apresentava cerca de trinta sprites de itens e NPCs sem código "
    "C# correspondente. A localização em português de Portugal estava parcialmente "
    "concluída: o Kraken encontrava-se traduzido, enquanto a maioria dos itens "
    "Monthra e Crystaline mantinha as entradas comentadas em inglês."
))

# ─── 3. REGISTO DE ALTERAÇÕES ──────────────────────────────────────────────
h1(doc, "3. Registo de Alterações")
np(doc, (
    "Esta secção regista cronologicamente todas as modificações efetuadas ao projeto "
    "após o início deste acompanhamento. Cada entrada inclui a data, o tipo de "
    "alteração (adição, modificação ou remoção), os ficheiros afetados e uma "
    "descrição técnica do trabalho realizado."
))
h2(doc, "3.1. Entradas Pendentes")
np(doc, (
    "Nenhuma alteração registada até ao momento. As entradas serão adicionadas à "
    "medida que o desenvolvimento progredir."
))

# ─── 4. REGISTO DE ERROS ───────────────────────────────────────────────────
h1(doc, "4. Registo de Erros e Resolução")
np(doc, (
    "Esta secção documenta todos os erros, falhas de compilação, crashes e "
    "comportamentos inesperados encontrados durante o desenvolvimento, bem como as "
    "respetivas resoluções aplicadas. Cada entrada indica a data, a descrição do "
    "erro, o contexto em que ocorreu e a solução implementada."
))
h2(doc, "4.1. Entradas Pendentes")
np(doc, "Nenhum erro registado até ao momento.")

# ─── 5. DECISÕES TÉCNICAS ──────────────────────────────────────────────────
h1(doc, "5. Decisões Técnicas e Notas de Desenvolvimento")
np(doc, (
    "Esta secção regista decisões de arquitetura e implementação que influenciam a "
    "direção do projeto, incluindo escolhas de design, opções descartadas e "
    "justificações técnicas relevantes."
))
h2(doc, "5.1. Entradas Pendentes")
np(doc, "Nenhuma decisão registada até ao momento.")

# ─── WEBGRAFIA ─────────────────────────────────────────────────────────────
h1(doc, "Webgrafia")
np(doc, (
    "tModLoader Team. tModLoader — Wiki e documentação oficial. "
    "Acedido em 11 de junho de 2026, em: https://github.com/tModLoader/tModLoader/wiki"
))
np(doc, (
    "Terraria Community Forums. "
    "Acedido em 11 de junho de 2026, em: https://forums.terraria.org/"
))

doc.save(OUTPUT)
print(f"Relatório gerado: {OUTPUT}")
